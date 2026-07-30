#!/usr/bin/env python3
"""
仓储突击检查SOP · Daryl修改版 → 模块B中越双语DOCX
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.1

BLUE = RGBColor(0x44, 0x72, 0xC4)
GRAY = RGBColor(0x88, 0x88, 0x88)
WARN = RGBColor(0xCC, 0x7A, 0x00)

def set_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def h(text, level=1):
    return doc.add_heading(text, level=level)

def p(text, bold=False, italic=False, size=None, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return para

def vn(text, size=10):
    """蓝色斜体越南语段落"""
    return p(text, italic=True, size=size, color=BLUE)

def warn(cn, vn_text=None):
    p(f'⚠ {cn}', size=10, color=WARN)
    if vn_text:
        vn(f'   {vn_text}')

def entry_block(lines_cn, lines_vn):
    for line in lines_cn:
        p2 = doc.add_paragraph()
        r = p2.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(10)
        p2.paragraph_format.space_after = Pt(1)
    sep = doc.add_paragraph()
    sr = sep.add_run('─' * 50)
    sr.font.size = Pt(7)
    sr.font.color.rgb = GRAY
    sep.paragraph_format.space_after = Pt(2)
    for line in lines_vn:
        p2 = doc.add_paragraph()
        r = p2.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(10)
        r.font.color.rgb = BLUE
        r.italic = True
        p2.paragraph_format.space_after = Pt(1)
    doc.add_paragraph()

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = hdr
        for pr in cell.paragraphs:
            for rn in pr.runs:
                rn.bold = True
                rn.font.size = Pt(9)
        set_shading(cell, 'E8E8E8')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for pr in cell.paragraphs:
                for rn in pr.runs:
                    rn.font.size = Pt(9)
    doc.add_paragraph()
    return table

# ═══════════ 标题 ═══════════
h('仓储突击检查SOP', level=0)
p('SOP Kiểm tra Đột xuất Kho hàng · 中越双语（模块B）', italic=True, size=10, color=GRAY)

# ═══════════ A · 盘点选样指引 ═══════════
h('A · 盘点选样指引——选样基础和规则', level=1)
p('核心逻辑：①抽验样本有效；②尽可能降低抽盘需要的时间；③抽验记录可追踪；')

h('A1', level=2)
p('原材料（纱线）、化工仓按环思：有效仓库位置+物料号10条筛选后，按物料号透视表的总数作为目标盘点数量；（见演示）')

h('A2', level=2)
p('坯布和成品仓：有效仓库位置+布号20条筛选后，根据库位前往仓库使用扫描枪扫码库位的所有存货，并在盘点表上记录数量作为扫描完整的依据；（见演示）')

h('A3', level=2)
p('五金仓和行政仓：BIP仓库实时数据导出随机抽取20条；')

# ═══════════════════════════════════════════════
# B · 盘点动作指引 — 中越双语 ★★★
# ═══════════════════════════════════════════════
h('B · 盘点动作指引——5模块问题发现标准动作', level=1)
vn('B · Hướng dẫn Thao tác Kiểm kê — 5 Module Phát hiện Vấn đề', size=13)

p('核心逻辑: 不是「去看看有没有问题」，而是每到一个仓库就按以下顺序走一遍。5个模块的顺序是固定的（由外到内、由表到账），形成肌肉记忆。')
vn('Logic cốt lõi: Không phải "đi xem có vấn đề gì không", mà là mỗi khi đến một kho, đi theo trình tự dưới đây. Thứ tự 5 module là cố định (từ ngoài vào trong, từ bề mặt đến sổ sách), tạo thành phản xạ cơ bắp.')

# ── B1 ──
h('B1. 分区定置标识（到仓第一眼——看分区、看标识、看危险品）', level=2)
vn('B1. Phân khu & Biển báo Định vị (Ấn tượng đầu tiên khi vào kho — Xem phân khu, biển báo, hàng nguy hiểm)')

p('检查逻辑: 站在仓库入口→沿通道走一遍→问题应该「扎眼」')
vn('Logic kiểm tra: Đứng ở cửa kho → Đi dọc lối đi một lượt → Vấn đề phải "đập vào mắt"')

make_table(
    ['动作编号\nMã HĐ', '检查动作 / Hành động kiểm tra', '怎么才算有问题 / Thế nào là có vấn đề', '拍照要求 / Yêu cầu chụp ảnh'],
    [
        ['B1-1',
         '看仓库入口有无平面分区图\nXem cửa kho có sơ đồ phân khu không',
         '无分区图，或有图但与实际货位不对应\nKhông có sơ đồ, hoặc có nhưng không khớp vị trí thực tế',
         '入口位置全景\nToàn cảnh vị trí cửa vào'],
        ['B1-2',
         '沿通道走，看货架/区域有无定置标识牌（如A区-五金、B区-油墨）\nĐi dọc lối đi, xem kệ/khu vực có biển định vị không (VD: Khu A-Ngũ kim, Khu B-Mực in)',
         '标识缺失、模糊、掉落未补\nBiển báo thiếu, mờ, rơi chưa gắn lại',
         '拍到无标识的区域全景+货架编号（如有）\nToàn cảnh khu vực thiếu biển + Mã kệ (nếu có)'],
        ['B1-3',
         '随机抽查3-5个货位，看物料卡/标签是否在位且信息完整（物料编码+名称+规格+单位）\nKiểm tra ngẫu nhiên 3-5 vị trí, xem thẻ/nhãn vật tư có đầy đủ thông tin (mã VT + tên + quy cách + ĐVT)',
         '无标签、标签信息不全、手写涂改未签字、标签与实物不符\nKhông nhãn, thiếu thông tin, viết tay sửa chưa ký, nhãn không khớp thực tế',
         '问题标签特写+所在货架全景\nCận cảnh nhãn lỗi + Toàn cảnh kệ'],
        ['B1-4',
         '关键项：扫一遍仓库，看有无危险品/易燃品（柴油、溶剂、化学品）与其他物料混放\nMục then chốt: Quét một lượt kho, xem hàng nguy hiểm/dễ cháy (dầu diesel, dung môi, hóa chất) có để lẫn với vật tư khác không',
         '①危化品无独立区域 ②无警示标识 ③无消防器材/消防通道被堵 ④化学品无二次防漏托盘\n①Hóa chất nguy hiểm không có khu riêng ②Không biển cảnh báo ③Không thiết bị PCCC/lối thoát hiểm bị chặn ④Hóa chất không có khay chống tràn thứ cấp',
         '危化品存放位置全景+警示标识缺失特写\nToàn cảnh nơi để hóa chất + Cận cảnh thiếu biển cảnh báo'],
        ['B1-5',
         '看通道是否畅通（主通道≥1.2m，次通道≥0.8m为参考线）\nXem lối đi có thông thoáng không (lối chính ≥1.2m, lối phụ ≥0.8m tham khảo)',
         '货物堆到通道、通道被叉车/托盘长时间占用\nHàng chất ra lối đi, lối đi bị xe nâng/pallet chiếm dụng lâu',
         '通道堵塞全景\nToàn cảnh lối đi bị tắc'],
    ]
)

p('常见问题速查:')
vn('Tra cứu nhanh vấn đề thường gặp:')
p('①柴油桶放五金仓角落无标识 ②油墨/稀释剂无防漏措施 ③货架标签用的是临时手写纸片')
vn('①Thùng dầu diesel để góc kho ngũ kim không biển báo ②Mực in/dung môi không biện pháp chống tràn ③Nhãn kệ dùng giấy viết tay tạm thời')

# ── B2 ──
h('B2. 现场周转控制（走动中观察——看状态、看日期、看呆滞）', level=2)
vn('B2. Kiểm soát Luân chuyển Hiện trường (Quan sát khi đi — Xem tình trạng, ngày tháng, hàng tồn đọng)')

p('检查逻辑: 边走边看物料的新旧程度、保质期、堆积时间')
vn('Logic kiểm tra: Vừa đi vừa xem mức độ cũ mới, hạn sử dụng, thời gian tồn đọng của vật tư')

make_table(
    ['动作编号\nMã HĐ', '检查动作 / Hành động kiểm tra', '怎么才算有问题 / Thế nào là có vấn đề', '拍照要求 / Yêu cầu chụp ảnh'],
    [
        ['B2-1',
         '看物料外观：有无明显过期标识（生产日期+保质期已过）、锈蚀/结块/变色/包装破损\nXem bề ngoài vật tư: Có nhãn quá hạn rõ ràng (NSX+HSD đã qua), rỉ sét/vón cục/đổi màu/rách bao bì',
         '目视即可判断的过期或变质\nQuá hạn hoặc biến chất có thể nhận biết bằng mắt thường',
         '过期品标签特写+整堆过期品全景\nCận cảnh nhãn hàng quá hạn + Toàn cảnh đống hàng quá hạn'],
        ['B2-2',
         '问保管员：呆滞品清单在哪里？多久更新一次？\nHỏi thủ kho: Danh sách hàng tồn đọng ở đâu? Bao lâu cập nhật một lần?',
         '①无呆滞品清单 ②有清单但超过3个月未更新 ③清单上的物料现场找不到或反过来\n①Không có DS hàng tồn đọng ②Có DS nhưng quá 3 tháng chưa cập nhật ③Vật tư trong DS không tìm thấy tại kho hoặc ngược lại',
         '呆滞品区域全景\nToàn cảnh khu hàng tồn đọng'],
        ['B2-3',
         '抽查3-5项呆滞品，看OA报废流程是否已完成\nKiểm tra 3-5 mục hàng tồn đọng, xem quy trình thanh lý OA đã hoàn tất chưa',
         'OA已批报废但实物仍在仓（说明处置脱节）\nOA đã duyệt thanh lý nhưng hàng vẫn còn trong kho (xử lý không đồng bộ)',
         'OA截图（保管员提供）+实物照片\nẢnh chụp OA (thủ kho cung cấp) + Ảnh hàng thực tế'],
        ['B2-4',
         '看退货区/待处理区的物料，问这些放了多久\nXem khu hàng trả lại/khu chờ xử lý, hỏi để bao lâu rồi',
         '超过1个月未处理的退货/待处理品且无跟进记录\nHàng trả lại/chờ xử lý quá 1 tháng không có hồ sơ theo dõi',
         '退货区全景+最近一次处理记录\nToàn cảnh khu trả hàng + Hồ sơ xử lý gần nhất'],
        ['B2-5',
         '问保管员近期入库/出库频率，到高频出入区域看地面有无散落、包装残损\nHỏi thủ kho tần suất nhập/xuất gần đây, đến khu vực có tần suất cao xem sàn có rơi vãi, bao bì hư hỏng',
         '地面有散落物料、破损包装未清理\nSàn có vật tư rơi vãi, bao bì rách chưa dọn',
         '散落区域全景\nToàn cảnh khu vực rơi vãi'],
    ]
)

p('常见问题速查:')
vn('Tra cứu nhanh vấn đề thường gặp:')
p('①化工类过期品OA已报废但实物未处置 ②退货区货物堆放无时间记录 ③呆滞品台账与实际对不上')
vn('①Hóa chất quá hạn OA đã thanh lý nhưng chưa xử lý thực tế ②Hàng khu trả lại chất đống không ghi thời gian ③Sổ hàng tồn đọng không khớp thực tế')

# ── B3 ──
h('B3. 收发存控制（到缓冲区/交接区——看交接、看管控、看记录）', level=2)
vn('B3. Kiểm soát Nhập-Xuất-Tồn (Đến khu đệm/khu bàn giao — Xem bàn giao, kiểm soát, hồ sơ)')

p('检查逻辑: 聚焦「货从哪来、怎么交接、没交接完的怎么管」')
vn('Logic kiểm tra: Tập trung "Hàng từ đâu đến, bàn giao thế nào, hàng chưa bàn giao xong quản lý ra sao"')

make_table(
    ['动作编号\nMã HĐ', '检查动作 / Hành động kiểm tra', '怎么才算有问题 / Thế nào là có vấn đề', '拍照要求 / Yêu cầu chụp ảnh'],
    [
        ['B3-1',
         '找到入库暂存区/缓冲区，看暂存货物有无临时标识（到货日期+供应商+物料名称）\nTìm khu tạm để nhập/khu đệm, xem hàng tạm có nhãn tạm thời (ngày đến + NCC + tên vật tư)',
         '暂存货物无任何标识，无法区分批次和归属\nHàng tạm không có bất kỳ nhãn nào, không phân biệt được lô và chủ sở hữu',
         '无标识暂存货全景\nToàn cảnh hàng tạm không nhãn'],
        ['B3-2',
         '问保管员：当日未入库的暂存货怎么管控？下班后谁负责？\nHỏi thủ kho: Hàng tạm chưa nhập trong ngày quản lý thế nào? Sau giờ làm ai chịu trách nhiệm?',
         '①无管控措施（无人锁门/无交接记录）②下班后缓冲区可自由进出\n①Không biện pháp kiểm soát (không khóa/không hồ sơ bàn giao) ②Sau giờ làm khu đệm có thể ra vào tự do',
         '缓冲区与外部通道连接处全景\nToàn cảnh chỗ nối khu đệm với lối đi bên ngoài'],
        ['B3-3',
         '看收发交接记录（送货单签收联/入库单），抽查最近3天\nXem hồ sơ bàn giao nhận (liên ký nhận phiếu giao hàng/phiếu nhập kho), kiểm tra 3 ngày gần nhất',
         '①无交接记录 ②交接记录无双方签字 ③到货日期与入库日期超过2个工作日且无说明\n①Không hồ sơ bàn giao ②Hồ sơ không có chữ ký hai bên ③Ngày đến hàng và ngày nhập kho quá 2 ngày làm việc không giải trình',
         '缺失的记录台账\nSổ sách/hồ sơ bị thiếu'],
        ['B3-4',
         '看出库是否有领料单，抽查最近3天的领料单\nXem xuất kho có phiếu lĩnh vật tư không, kiểm tra 3 ngày gần nhất',
         '①无领料单 ②领料单无审批签字 ③领用数量与系统扣减数量不一致\n①Không phiếu lĩnh ②Phiếu lĩnh không chữ ký phê duyệt ③SL lĩnh không khớp SL trừ hệ thống',
         '问题单据特写\nCận cảnh chứng từ có vấn đề'],
        ['B3-5',
         '问保管员：紧急领料（无单先领）怎么处理？补单时限？\nHỏi thủ kho: Lĩnh vật tư khẩn cấp (không phiếu, lấy trước) xử lý thế nào? Thời hạn bổ sung phiếu?',
         '①无制度 ②有制度但未执行（抽查到无单领料且超期未补）\n①Không quy định ②Có quy định nhưng không thực hiện (phát hiện lĩnh không phiếu quá hạn chưa bổ sung)',
         '相关记录\nHồ sơ liên quan'],
    ]
)

p('常见问题速查:')
vn('Tra cứu nhanh vấn đề thường gặp:')
p('①缓冲区夜间无管控 ②送货单上只有供应商签字无仓库签字 ③无单领料后长期不补单')
vn('①Khu đệm ban đêm không kiểm soát ②Phiếu giao hàng chỉ có chữ ký NCC, không có chữ ký kho ③Lĩnh không phiếu lâu ngày không bổ sung')

# ── B4 ──
h('B4. 仓库数据维护（坐到保管员电脑前——看系统、看单据、看时效）', level=2)
vn('B4. Duy trì Dữ liệu Kho (Ngồi trước máy tính thủ kho — Xem hệ thống, chứng từ, tính kịp thời)')

p('检查逻辑: 账实相符的前提是数据录入及时。不看全部数据，只抽最近几天的。')
vn('Logic kiểm tra: Tiền đề khớp đúng sổ sách-thực tế là dữ liệu được nhập kịp thời. Không xem toàn bộ, chỉ kiểm tra mấy ngày gần đây.')

make_table(
    ['动作编号\nMã HĐ', '检查动作 / Hành động kiểm tra', '怎么才算有问题 / Thế nào là có vấn đề', '拍照要求 / Yêu cầu chụp ảnh'],
    [
        ['B4-1',
         '打开仓库入库单台账，抽查最近一周\nMở sổ phiếu nhập kho, kiểm tra 1 tuần gần nhất',
         '①无台账 ②台账不是连续的（有跳跃编号）③入库单与台账日期不一致\n①Không sổ ②Sổ không liên tục (có số nhảy) ③Phiếu nhập và sổ không khớp ngày',
         '台账翻阅全景\nToàn cảnh sổ sách'],
        ['B4-2',
         '打开仓库出库/领料单台账，抽查最近一周\nMở sổ phiếu xuất kho/lĩnh vật tư, kiểm tra 1 tuần gần nhất',
         '①无台账 ②领料单未在次日12:00前录入系统（或已录入但无签名确认）\n①Không sổ ②Phiếu lĩnh chưa nhập hệ thống trước 12h ngày hôm sau (hoặc đã nhập nhưng chưa ký xác nhận)',
         '抽查到的未录入单据特写\nCận cảnh chứng từ chưa nhập'],
        ['B4-3',
         '抽查3-5张最近入库/出库单，反向追踪到系统（BIP/ERP），看系统数据与单据是否一致\nKiểm tra 3-5 phiếu nhập/xuất gần nhất, truy ngược lên hệ thống (BIP/ERP), xem dữ liệu có khớp chứng từ không',
         '单据有、系统无；或系统有、单据找不到\nCó chứng từ, hệ thống không có; hoặc hệ thống có, không tìm thấy chứng từ',
         '单据vs系统截图对比\nĐối chiếu ảnh chụp chứng từ vs hệ thống'],
        ['B4-4',
         '问保管员：自盘记录在哪里？最近一次什么时候？\nHỏi thủ kho: Hồ sơ tự kiểm kê ở đâu? Lần gần nhất khi nào?',
         '①无自盘记录 ②最近一次超过1个月 ③自盘差异未跟进处理\n①Không hồ sơ tự KK ②Lần gần nhất quá 1 tháng ③Chênh lệch tự KK chưa xử lý',
         '自盘记录本翻阅\nSổ tự kiểm kê'],
        ['B4-5',
         '问保管员：重要内控报告（如月末库存报表、差异分析）归档在哪里\nHỏi thủ kho: Các báo cáo kiểm soát nội bộ quan trọng (BC tồn kho cuối tháng, phân tích chênh lệch) lưu ở đâu',
         '①无归档 ②找不到 ③报告有但缺少当月/当季的\n①Không lưu ②Không tìm thấy ③Có BC nhưng thiếu tháng/quý hiện tại',
         '归档位置\nVị trí lưu trữ'],
    ]
)

p('常见问题速查:')
vn('Tra cứu nhanh vấn đề thường gặp:')
p('①入库频率低的仓库减少录入频次导致系统延迟 ②单据归档混乱（入库单和出库单混放、缺少日期排序）③自盘记录缺失')
vn('①Kho tần suất nhập thấp giảm số lần nhập dẫn đến hệ thống trễ ②Chứng từ lưu lộn xộn (phiếu nhập xuất để lẫn, thiếu sắp xếp theo ngày) ③Thiếu hồ sơ tự kiểm kê')

# ── B5 ──
h('B5. 盘点管理（回到盘点本身——看配合、看流程、看复盘）', level=2)
vn('B5. Quản lý Kiểm kê (Quay lại bản thân việc kiểm kê — Xem phối hợp, quy trình, rút kinh nghiệm)')

p('检查逻辑: 这部分是「对这次检查本身的复盘」，也是下次改进的依据')
vn('Logic kiểm tra: Phần này là "rút kinh nghiệm cho chính đợt kiểm tra này", cũng là cơ sở cải tiến cho lần sau')

make_table(
    ['动作编号\nMã HĐ', '检查动作 / Hành động kiểm tra', '怎么才算有问题 / Thế nào là có vấn đề', '拍照要求 / Yêu cầu chụp ảnh'],
    [
        ['B5-1',
         '记录本次抽盘实际耗时\nGhi nhận thời gian thực tế kiểm kê đợt này',
         '超过1.5小时→下次选料号要减少；不足0.5小时→下次可增加\nQuá 1.5h→Lần sau giảm số mã; Dưới 0.5h→Lần sau có thể tăng',
         '—'],
        ['B5-2',
         '记录仓储部配合情况：是否指定了带路人？扫码枪/叉车/登高设备是否可用？物料是否可触及？\nGhi nhận tình hình phối hợp của bộ phận kho: Có chỉ định người dẫn đường? Súng quét/xe nâng/thiết bị lên cao có hoạt động? Vật tư có thể tiếp cận?',
         '①无人带路 ②设备故障影响盘点 ③高位物料无法安全取用\n①Không người dẫn ②Thiết bị hỏng ảnh hưởng KK ③Hàng trên cao không thể lấy an toàn',
         '设备故障/无法盘点位置\nThiết bị hỏng / Vị trí không thể KK'],
        ['B5-3',
         '记录盘点中断原因（如有）\nGhi nhận nguyên nhân gián đoạn KK (nếu có)',
         '任何导致抽盘未能100%完成的原因\nBất kỳ nguyên nhân nào khiến KK không hoàn thành 100%',
         '—'],
        ['B5-4',
         '盘点结束后与保管员当面确认差异，双方签字\nSau KK, xác nhận chênh lệch trực tiếp với thủ kho, hai bên ký tên',
         '保管员拒绝签字→记入报告「配合异常」\nThủ kho từ chối ký→Ghi vào báo cáo "Phối hợp bất thường"',
         '签字后的盘点表\nBiểu KK đã ký'],
        ['B5-5',
         '盘点结束后组内快速复盘（≤5分钟），确认本次检查的3个最值得记录的点（问题/亮点/改进建议各1）\nSau KK, nhóm rút kinh nghiệm nhanh (≤5 phút), xác nhận 3 điểm đáng ghi nhận nhất (1 vấn đề/1 điểm sáng/1 đề xuất cải tiến)',
         '—',
         '—'],
    ]
)

p('常见问题速查:')
vn('Tra cứu nhanh vấn đề thường gặp:')
p('①财务选料号太多导致超时 ②扫码枪没电/故障/不兼容 ③高位货架物料无法清点')
vn('①Kế toán chọn quá nhiều mã dẫn đến quá giờ ②Súng quét hết pin/hỏng/không tương thích ③Hàng kệ cao không thể kiểm đếm')

# ── B 动作检查顺序总结 ──
h('动作检查顺序总结（肌肉记忆版）', level=3)
vn('Tổng kết Trình tự Thao tác Kiểm tra (Phiên bản Phản xạ Cơ bắp)')

entry_block(
    ['到仓 → B1 走一圈看标识和危险品（5分钟）',
     ' → B2 边走边看呆滞过期（5分钟）',
     ' → B3 到缓冲区看交接管控（5分钟）',
     ' → B4 坐到电脑前抽单据台账（10分钟）',
     ' → B5 盘点结束记配合度+复盘（最后5分钟）',
     '',
     '总耗时控制：30分钟/仓（不含物料逐一清点时间）',
     '物料清点时间：按抽样数量另计'],
    ['Đến kho → B1 Đi một vòng xem biển báo & hàng nguy hiểm (5 phút)',
     ' → B2 Vừa đi vừa xem hàng tồn đọng & quá hạn (5 phút)',
     ' → B3 Đến khu đệm xem bàn giao & kiểm soát (5 phút)',
     ' → B4 Ngồi trước máy tính kiểm tra chứng từ & sổ sách (10 phút)',
     ' → B5 Kết thúc KK ghi nhận mức độ phối hợp & rút kinh nghiệm (5 phút cuối)',
     '',
     'Tổng thời gian: 30 phút/kho (không bao gồm thời gian kiểm đếm từng mã)',
     'Thời gian kiểm đếm: Tính riêng theo số lượng mẫu']
)

# ═══════════════════════════════════════════════
# C · 盘点报告与跟进 (保持中文)
# ═══════════════════════════════════════════════
h('C · 盘点报告汇报与整改进度跟进', level=1)
p('核心原则: 检查完不汇报=没检查。汇报不走闭环=白检查。')

h('C1. 汇报架构（谁向谁汇报什么）', level=2)
p('越南财务部(执行层) / 中国总部(监督层)')
p('├─ 检查组长 ──→ 邮件报告 ──→ 仓储部负责人，事业部财务负责人', size=10)
p('│ + 抄送: 副总裁：林金建', size=10)
p('│ 整改清单 + 升级抄送: 事业部总经理', size=10)
p('├─ 检查组长 ──→ 口头简报 ──→ 仓库负责人（当场）', size=10)
p('└─ 检查组长 ──→ 整改跟踪表 ──→ 共享文档（持续更新）', size=10)

h('C2. 汇报节奏', level=2)
make_table(
    ['节点', '动作', '时限', '产出方式'],
    [
        ['当场', '口头简报仓库负责人', '盘点结束前', '口头'],
        ['24h内', '组员录入盘点数据到共享文档', '次日下班前', '共享文档'],
        ['48h内', '组长发邮件报告+整改清单', '第2个工作日17:00前', '邮件+附件'],
        ['持续', '跟进整改状态，更新跟踪表', '实时', '共享文档'],
        ['下期检查前', '对照上期清单逐项复核', '检查前1天', '复核记录'],
    ]
)

h('C3. 邮件报告模板', level=2)
p('参照Daryl的B类4段式结构：事由→人做什么→结果→未决+证据')

entry_block(
    ['主题：【仓储突击检查】[仓库名称] 第[ ]期 [YYYY-MM-DD]',
     '收件人：仓储部负责人、事业部财务负责人',
     '抄送：林金建、[事业部负责人]、盘点组员、会计长、仓储课长',
     '',
     '各位领导同事：',
     '一、事由与执行',
     '接仓储专项检查工作计划，财务部[姓名]+[姓名]组成第[X]突击检查小组，',
     '于[YYYY年MM月DD日 HH:MM-HH:MM]对[仓库名称]进行突击检查。',
     '本次抽查[X]个料号，账实相符[X]项，差异[X]项。',
     '',
     '二、发现问题（按5模块检查结果汇总）',
     '三、未完事项',
     '四、附件：盘点表 + 整改清单 + 检查照片'],
    ['(Giữ nguyên mẫu tiếng Trung — mẫu email báo cáo của Daryl)']
)

h('C4. 整改跟踪表', level=2)
p('建议用企微多维表格或共享Excel维护，字段：检查日期/期数/仓库/保管人/问题编号/模块/严重级别/问题描述/整改措施/责任人/承诺完成日/实际完成日/整改证据/财务复核/升级标记')

h('C5. 升级规则', level=2)
make_table(
    ['触发条件', '升级到', '升级方式'],
    [
        ['🔴项超承诺期限3天未完成', '事业部财务负责人', '组长口头+邮件提醒'],
        ['🔴项超承诺期限7天未完成', '事业部总经理', '事业部财务负责人邮件升级'],
        ['同一仓库连续2期出现同类🔴项', '事业部总经理', '财务部在报告中标注「重复发生」'],
        ['同一仓库连续3期出现同类🔴项', '集团CFO', '事业部财务负责人邮件升级'],
        ['仓储部拒绝配合检查', '事业部总经理', '组长当场电话+事后邮件'],
    ]
)

# ── 保存 ──
output = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'reports', '仓储突击检查SOP-B部分-中越双语-20260730.docx')
output = os.path.abspath(output)
os.makedirs(os.path.dirname(output), exist_ok=True)
doc.save(output)
print(f'✅ 已保存: {output}')
