#!/usr/bin/env python3
"""
将应付采购入账SOP v5.0 中越双语 HTML → DOCX
保持原docx结构风格，输出为标准docx格式
"""
import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ═══ 样式设置 ═══
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bilingual_para(cn_text, vn_text, cn_bold=False):
    """添加中越双语段落"""
    p = doc.add_paragraph()
    run_cn = p.add_run(cn_text)
    run_cn.bold = cn_bold
    p2 = doc.add_paragraph()
    run_vn = p2.add_run(vn_text)
    run_vn.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)  # 蓝色
    run_vn.font.size = Pt(10)
    run_vn.italic = True
    p2.paragraph_format.space_after = Pt(2)
    return p, p2

def add_entry_block(lines_cn, lines_vn):
    """添加分录块（中越对照）"""
    for line in lines_cn:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
    # 分隔
    sep = doc.add_paragraph()
    sep_run = sep.add_run('─' * 50)
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    sep.paragraph_format.space_after = Pt(2)
    sep.paragraph_format.space_before = Pt(2)
    for line in lines_vn:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run.italic = True
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
    # 空行
    doc.add_paragraph()

def add_warning(text_cn, text_vn=None):
    p = doc.add_paragraph()
    run = p.add_run('⚠ ' + text_cn)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
    if text_vn:
        p2 = doc.add_paragraph()
        run2 = p2.add_run('   ' + text_vn)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run2.italic = True

def make_table(headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        set_cell_shading(cell, 'E8E8E8')
    
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()  # spacing
    return table


# ═══════════════════════════════════════════════
# 文档主体
# ═══════════════════════════════════════════════

add_heading('应付采购入账SOP — 框架 v5.0', level=0)
add_para('SOP Hạch toán Phải trả Người bán · Phiên bản 5.0 · 2026-07-30 · 中越双语（第三章）', 
         italic=True, size=10, color=(0x88,0x88,0x88))

# ── 阅读导航 ──
add_heading('📑 阅读导航 / Hướng dẫn đọc', level=2)
make_table(
    ['角色 Vai trò', '必读章节 Chương cần đọc', '目的 Mục đích'],
    [
        ['财务·日常入账', '第1章 + 第2章 + 第3章 + 第4章', '全覆盖，按确认→计量→记录→期末处理的会计逻辑推进'],
        ['财务·新人培训', '全文（跳过附录）', '建立完整认知框架'],
        ['采购部', '1.1(确认时点) + 3.3(单据清单) + 3.4(单据流转)', '知道什么时候该交什么单据'],
        ['仓库', '1.1(收货) + 3.3(入库单等)', '收货即出单，不压单'],
        ['内控/审计', '1.1 + 2.2 + 3.3 + 4.5 + 4.6', '控制点检查'],
        ['管理层', '阅读导航 + 附录B', '拍板用'],
    ]
)

# ═══════════════════════════════════════════════
# 第一章
# ═══════════════════════════════════════════════
add_heading('第一章 · 确认 — Recognition', level=1)

add_heading('1.1 入账时点认定', level=2)
add_para('控制权转移的判断依据两个要素：①我有权支配这个资产 ②我承担了资产的风险和报酬。')

add_heading('本地采购', level=3)
make_table(
    ['场景', '控制权转移时点', '入账日', '依据'],
    [
        ['标准本地采购', '收货验收完成 + 收到VAT发票', '收货验收日（发票后补走暂估）', 'VAS 02'],
        ['合同约定验收期', '验收期届满或提前验收通过', '验收通过日', '合同条款 > 默认规则'],
        ['分期交付', '每批独立验收', '各批验收日', '分批控制权转移'],
    ]
)
add_warning('实务铁律：收货验收完成 = 控制权转移，不以发票收到为入账前提。发票未到走暂估（见4.1）。')

add_heading('进口采购', level=3)
make_table(
    ['Incoterms', '控制权转移时点', '实务入账日', '说明'],
    [
        ['FOB', '装船越过船舷', '报关放行日', '货款金额需报关才能确定'],
        ['CIF', '货物在目的港越过船舷', '报关放行日', '卖方承担运费+保险'],
        ['DAP/DDP', '货物到达买方指定地点', '到厂验收日', '卖方承担全部运输风险'],
        ['FCA', '货交第一承运人', '报关放行日', '等同FOB的陆运/空运版'],
        ['EXW', '卖方工厂交货', '提货日', '买方承担全部运输风险'],
    ]
)
add_warning('SOP 统一规则：进口采购统一以报关放行日为入账时点。理由：①税款金额此时确定 ②海关放行意味着有合法支配权 ③实操中这是最早可操作的时点。')

add_heading('委外加工（特殊）', level=3)
make_table(
    ['阶段', '确认什么', '时点'],
    [
        ['发出材料', '不确认损益，仅资产内部转移', '材料出库日'],
        ['加工费', '收到加工方发票', '发票日'],
        ['收回成品', '成本结转', '成品入库日'],
    ]
)

add_heading('固定资产采购', level=3)
make_table(
    ['类型', '入账时点', '折旧起点'],
    [
        ['不需要安装', '验收合格日', '投入使用次月'],
        ['需要安装', '达到预定可使用状态日', '投入使用次月'],
    ]
)

add_heading('1.2 关键判断：Incoterms与在途物资', level=2)
make_table(
    ['判断', '规则'],
    [
        ['FOB已装船、未到港', '实务中多数等报关放行日入账。SOP默认不挂账，等报关日。'],
        ['DAP已到港、尚未到厂', '控制权在到厂时才转移，等验收日。'],
        ['跨月报关', '月末最后一天放行 → 必须当月入账，不能跨月'],
    ]
)

# ═══════════════════════════════════════════════
# 第二章
# ═══════════════════════════════════════════════
add_heading('第二章 · 计量 — Measurement', level=1)

add_heading('2.1 采购成本构成通用框架', level=2)
add_para('存货/固资入账成本 = 采购价格（不含VAT）+ 直接归属费用（使资产达到预定状态/地点的必要支出）- 采购折扣/折让')

add_heading('2.2 成本构成按类展开', level=2)

add_heading('2.2.1 原材料 (152)', level=3)
make_table(
    ['成本项', '本地', '进口'],
    [
        ['不含税买价', '✅ 发票金额/(1+VAT%)', '✅ CIF价(VND换算后)'],
        ['运费', '✅ 合同约定自提的运费', '✅ 海运费'],
        ['保险费', '—', '✅ 运输险'],
        ['进口关税(不可退)', '—', '✅ 计入成本'],
        ['进口关税(可退)', '—', '❌ 挂3333等退税'],
        ['港杂/拖车/报关代理', '—', '✅ 计入成本'],
    ]
)

add_heading('2.2.2 生物质材料 / 2.2.3 化学品 / 2.2.4 委外加工 / 2.2.5 固定资产', level=3)
add_para('（详细内容见原版文档，此处保持原有格式不变）')

add_heading('2.3 进口附加：关税/汇率/运费分摊', level=2)
add_para('详见原版文档 2.3.1-2.3.3')

add_heading('2.4 价外成本归集', level=2)
add_para('详见原版文档')

add_heading('2.5 采购折让与折扣', level=2)
add_warning('铁律：采购折扣/折让一律冲减采购成本，不进其他业务收入(711)')

# ═══════════════════════════════════════════════
# 第三章 · 中越双语 ★★★
# ═══════════════════════════════════════════════
add_heading('第三章 · 入账 — Accounting', level=1)
add_para('Chương 3 · Hạch toán', italic=True, size=14, color=(0x44,0x72,0xC4))
add_para('确认和计量完成后，入账是落地环节。', bold=False)
add_para('Sau khi hoàn tất ghi nhận và đo lường, hạch toán là bước hiện thực hóa vào sổ sách.', italic=True, size=10, color=(0x44,0x72,0xC4))

# ── 3.1 标准分录 ──
add_heading('3.1 标准分录', level=2)
add_para('3.1 Bút toán chuẩn', italic=True, size=13, color=(0x44,0x72,0xC4))

# 3.1.1
add_heading('3.1.1 原材料·本地采购', level=3)
add_para('3.1.1 Nguyên vật liệu · Mua trong nước', italic=True, size=11, color=(0x44,0x72,0xC4))

add_entry_block(
    ['入账:',
     '  借: 152 - Nguyên liệu, vật liệu  [不含税价]',
     '  借: 1331 - Thuế GTGT được khấu trừ  [VAT]',
     '  贷: 331 - Phải trả người bán  [含税总价]',
     '如有付现运费:',
     '  借: 152    借: 1331 (如有)    贷: 111/112',
     '质量折让(不退货):',
     '  借: 331 [折让金额含税]    贷: 152 [折让不含税]    贷: 1331 [折让VAT]',
     '付款:',
     '  借: 331    贷: 112'],
    ['Hạch toán:',
     '  Nợ: 152 - Nguyên liệu, vật liệu  [Giá chưa VAT]',
     '  Nợ: 1331 - Thuế GTGT được khấu trừ  [VAT]',
     '  Có: 331 - Phải trả người bán  [Tổng giá có VAT]',
     'Nếu có cước vận chuyển trả ngay:',
     '  Nợ: 152    Nợ: 1331 (nếu có)    Có: 111/112',
     'Giảm giá chất lượng (không trả hàng):',
     '  Nợ: 331 [Số giảm giá có VAT]    Có: 152 [Chưa VAT]    Có: 1331 [VAT]',
     'Thanh toán:',
     '  Nợ: 331    Có: 112']
)

# 3.1.2
add_heading('3.1.2 原材料·进口采购', level=3)
add_para('3.1.2 Nguyên vật liệu · Nhập khẩu', italic=True, size=11, color=(0x44,0x72,0xC4))

add_entry_block(
    ['入账:',
     '  借: 152  [CIF货价+关税(不可退)+港杂]',
     '  借: 1331  [进口VAT = (CIF+关税)×VAT%]',
     '  贷: 331 - (境外供应商)  [CIF货价(VND)]',
     '  贷: 3333 - Thuế nhập khẩu  [进口关税]',
     '  贷: 33312 - Thuế GTGT hàng nhập khẩu  [进口VAT]',
     '支付关税+VAT:',
     '  借: 3333    借: 33312    贷: 112',
     '支付付现费用(拖车/港杂):',
     '  借: 152    借: 1331 (如有)    贷: 111/112',
     '支付供应商货款:',
     '  借: 331    贷: 112',
     '  → 汇率差异: Dr 635 或 Cr 515'],
    ['Hạch toán:',
     '  Nợ: 152  [Giá CIF + Thuế NK (không hoàn) + Phí cảng]',
     '  Nợ: 1331  [VAT NK = (CIF + Thuế NK) × %VAT]',
     '  Có: 331 - (NCC nước ngoài)  [Giá CIF quy VND]',
     '  Có: 3333 - Thuế nhập khẩu',
     '  Có: 33312 - Thuế GTGT hàng nhập khẩu',
     'Nộp thuế NK + VAT:',
     '  Nợ: 3333    Nợ: 33312    Có: 112',
     'Thanh toán phí kéo container, phí cảng:',
     '  Nợ: 152    Nợ: 1331 (nếu có)    Có: 111/112',
     'Thanh toán tiền hàng cho NCC:',
     '  Nợ: 331    Có: 112',
     '  → Chênh lệch tỷ giá: Nợ 635 hoặc Có 515']
)

# 3.1.3
add_heading('3.1.3 生物质材料·本地采购', level=3)
add_para('3.1.3 Nguyên liệu sinh khối · Mua trong nước', italic=True, size=11, color=(0x44,0x72,0xC4))
add_para('核心分录同原材料(3.1.1) + 含水率调整。过磅后含水率高于基准 → 结算重量调减：')
add_para('Bút toán cơ bản như NVL (3.1.1) + Điều chỉnh độ ẩm. Độ ẩm thực tế cao hơn chuẩn → Giảm khối lượng thanh toán:', italic=True, size=10, color=(0x44,0x72,0xC4))
add_entry_block(
    ['  借: 152 [调减后不含税价]    借: 1331    贷: 331 [调减后含税价]',
     '  结算重量 = 过磅重量 × (1 - 实际含水率) / (1 - 基准含水率)'],
    ['  Nợ: 152 [Giá chưa VAT sau điều chỉnh]    Nợ: 1331    Có: 331 [Giá có VAT]',
     '  Khối lượng thanh toán = KL cân × (1 - Độ ẩm thực tế) / (1 - Độ ẩm chuẩn)']
)

# 3.1.4
add_heading('3.1.4 化学品·本地采购', level=3)
add_para('3.1.4 Hóa chất · Mua trong nước', italic=True, size=11, color=(0x44,0x72,0xC4))
add_para('核心分录同原材料(3.1.1) + 环保税：')
add_para('Bút toán cơ bản như NVL (3.1.1) + Thuế bảo vệ môi trường:', italic=True, size=10, color=(0x44,0x72,0xC4))
add_entry_block(
    ['缴纳环保税:',
     '  借: 152 [环保税金额]    贷: 3338 - Thuế bảo vệ môi trường',
     '  借: 3338    贷: 112'],
    ['Nộp thuế BVMT:',
     '  Nợ: 152 [Số thuế BVMT]    Có: 3338 - Thuế bảo vệ môi trường',
     '  Nợ: 3338    Có: 112']
)

# 3.1.5
add_heading('3.1.5 委外加工', level=3)
add_para('3.1.5 Gia công bên ngoài', italic=True, size=11, color=(0x44,0x72,0xC4))

add_entry_block(
    ['① 发出材料:',
     '  借: 154 - CPSXKD dở dang (bên ngoài GC)    贷: 152',
     '② 加工费:',
     '  借: 154 [加工费不含税]    借: 1331    贷: 331 - (加工方)',
     '③ 加工运杂费:',
     '  借: 154    借: 1331 (如有)    贷: 111/112/331',
     '④ 收回成品:',
     '  借: 152(半成品) / 155(产成品)    贷: 154',
     '⑤ 余料退回:',
     '  借: 152    贷: 154',
     '⑥ 超耗索赔:',
     '  借: 331 / 1388 [加工方承担]    贷: 154'],
    ['① Xuất NVL đi gia công:',
     '  Nợ: 154 - CPSXKD dở dang (bên ngoài GC)    Có: 152',
     '② Phí gia công:',
     '  Nợ: 154 [Phí GC chưa VAT]    Nợ: 1331    Có: 331 - (Bên GC)',
     '③ Cước vận chuyển GC:',
     '  Nợ: 154    Nợ: 1331 (nếu có)    Có: 111/112/331',
     '④ Nhập lại thành phẩm:',
     '  Nợ: 152(BTP) / 155(Thành phẩm)    Có: 154',
     '⑤ NVL thừa trả về:',
     '  Nợ: 152    Có: 154',
     '⑥ Hao hụt vượt mức – yêu cầu bồi thường:',
     '  Nợ: 331 / 1388 [Bên GC chịu]    Có: 154']
)

# 3.1.6
add_heading('3.1.6 固定资产·本地采购', level=3)
add_para('3.1.6 TSCĐ · Mua trong nước', italic=True, size=11, color=(0x44,0x72,0xC4))

add_entry_block(
    ['入账:',
     '  借: 211 - TSCĐ hữu hình  [不含税价+运杂+安装调试]',
     '  借: 1332 - Thuế GTGT được KT (TSCĐ)  [固资VAT单独核算]',
     '  贷: 331  [含税总价]'],
    ['Hạch toán:',
     '  Nợ: 211 - TSCĐ hữu hình  [Giá chưa VAT + Vận chuyển + Lắp đặt]',
     '  Nợ: 1332 - Thuế GTGT được KT (TSCĐ)  [VAT TSCĐ riêng]',
     '  Có: 331  [Tổng giá có VAT]']
)

# 3.1.7
add_heading('3.1.7 固定资产·进口采购', level=3)
add_para('3.1.7 TSCĐ · Nhập khẩu', italic=True, size=11, color=(0x44,0x72,0xC4))

add_entry_block(
    ['需要安装 → 先在在建工程归集:',
     '  借: 241 - XDCB dở dang  [CIF+关税(不可退)+港杂+安装费]',
     '  借: 1332    贷: 331    贷: 3333    贷: 33312',
     '达到可使用状态 → 转固:',
     '  借: 211    贷: 241',
     '次月折旧:',
     '  借: 627/642 (按使用部门)    贷: 214 - Khấu hao TSCĐ'],
    ['Cần lắp đặt → Tập hợp vào XDCB dở dang:',
     '  Nợ: 241 - XDCB dở dang  [CIF + Thuế NK (không hoàn) + Phí cảng + Lắp đặt]',
     '  Nợ: 1332    Có: 331    Có: 3333    Có: 33312',
     'Đạt trạng thái sẵn sàng sử dụng → Chuyển TSCĐ:',
     '  Nợ: 211    Có: 241',
     'Trích khấu hao tháng kế tiếp:',
     '  Nợ: 627/642 (theo bộ phận)    Có: 214 - Khấu hao TSCĐ']
)

# ── 3.2 科目映射表 ──
add_heading('3.2 科目映射表', level=2)
add_para('3.2 Bảng ánh xạ tài khoản', italic=True, size=13, color=(0x44,0x72,0xC4))

make_table(
    ['业务 / Nghiệp vụ', '资产科目 TK TS', '进项税 TK VAT', '负债科目 TK Nợ', '暂挂 TK tạm'],
    [
        ['原材料·本地 / NVL Trong nước', '152', '1331', '331', '—'],
        ['原材料·进口 / NVL Nhập khẩu', '152', '1331', '331+3333+33312', '—'],
        ['生物质·本地 / Sinh khối', '152', '1331', '331', '—'],
        ['化学品·本地 / Hóa chất', '152', '1331', '331+3338', '—'],
        ['委外加工 / Gia công ngoài', '154', '1331', '331', '—'],
        ['固资·本地 / TSCĐ Trong nước', '211', '1332', '331', '—'],
        ['固资·进口 / TSCĐ Nhập khẩu', '211/241', '1332', '331+3333+33312', '—'],
        ['暂估 / Tạm tính', '152/211', '—', '338-暂估', '—'],
        ['预付 / Trả trước', '—', '—', '331-预付', 'Dr 331'],
        ['汇率重估 / Đánh giá tỷ giá', '—', '—', '331', 'Dr 635/Cr 515'],
    ]
)

add_warning(
    '1331 vs 1332：货物和服务的进项税用1331，固定资产的进项税用1332。两者都可以抵扣，但分开核算（税务申报表也分开填）。',
    '1331 vs 1332：VAT đầu vào của hàng hóa, dịch vụ dùng TK 1331; VAT đầu vào của TSCĐ dùng TK 1332. Cả hai đều được khấu trừ nhưng hạch toán riêng (tờ khai thuế cũng kê riêng).'
)

# ── 3.3 单据完整性清单 ──
add_heading('3.3 单据完整性清单', level=2)
add_para('3.3 Danh sách chứng từ đầy đủ', italic=True, size=13, color=(0x44,0x72,0xC4))

add_heading('通用单据 / Chứng từ chung', level=3)
make_table(
    ['级别', '单据 / Chứng từ', '来源 / Nguồn', '缺失后果 / Hậu quả nếu thiếu'],
    [
        ['🔴', '采购订单(PO) / Đơn đặt hàng', '采购部 / P.Mua hàng', '无法证明授权 / Không CM được phê duyệt'],
        ['🔴', '入库单/收货单 / Phiếu nhập kho', '仓库 / Kho', '无法确认入库 / Không XN thực nhập'],
        ['🔴', 'VAT发票 / Hóa đơn GTGT', '供应商 / NCC', '无法抵扣进项税 / Không được KT VAT'],
        ['🔴', '采购合同(关键条款) / Hợp đồng', '采购部 / P.Mua hàng', '付款条件无依据 / Không căn cứ TT'],
        ['🟠', '供应商送货单 / Phiếu giao hàng NCC', '仓库 / Kho', '收发差异无法追溯'],
        ['🟠', '质检报告 / BB kiểm tra chất lượng', '质检部 / QC', '质量无保证'],
        ['🟢', '请购单(PR) / Phiếu yêu cầu mua', '需求部门 / BP yêu cầu', '审批留痕'],
        ['🟢', '比价记录 / Bảng so sánh giá', '采购部 / P.Mua hàng', '内控审计'],
    ]
)

add_heading('进口附加单据 / Chứng từ bổ sung – Nhập khẩu', level=3)
make_table(
    ['级别', '单据 / Chứng từ', '来源 / Nguồn'],
    [
        ['🔴', '商业发票 / Commercial Invoice', '境外供应商 / NCC nước ngoài'],
        ['🔴', '装箱单 / Packing List', '境外供应商 / NCC nước ngoài'],
        ['🔴', '提单 / Bill of Lading', '船公司 / Hãng tàu'],
        ['🔴', '报关单 / Tờ khai hải quan', '海关/报关行 / HQ/Đại lý HQ'],
        ['🔴', '关税缴款书 / Biên lai thuế NK', '海关 / Hải quan'],
        ['🔴', '进口VAT缴款书 / Biên lai VAT NK', '海关 / Hải quan'],
        ['🟠', '运费发票+保险单 / Hóa đơn cước + BH', '货代/保险 / Forwarder/BH'],
        ['🟠', '港杂/拖车费凭证 / Chứng từ phí cảng', '物流 / Logistics'],
        ['🟢', '报关委托书 / Giấy ủy quyền khai HQ', '报关行 / Đại lý HQ'],
    ]
)

add_heading('委外加工单据 / Chứng từ – Gia công ngoài', level=3)
make_table(
    ['级别', '单据 / Chứng từ', '来源 / Nguồn'],
    [
        ['🔴', '委外加工合同 / Hợp đồng gia công', '采购部 / P.Mua hàng'],
        ['🔴', '材料出库单(发外) / Phiếu xuất NVL đi GC', '仓库 / Kho'],
        ['🔴', '加工完成入库单 / Phiếu nhập kho TP GC', '仓库 / Kho'],
        ['🟠', '损耗报告 / Báo cáo hao hụt', '加工方 / Bên GC'],
        ['🟠', '替代：采购部统计台账和委外订单分析报告\n      Sổ theo dõi TK & BC phân tích đơn GC', '采购部'],
        ['🟠', '余料退回单 / Phiếu nhập NVL thừa', '仓库 / Kho'],
        ['🟠', '月度盘点表（委外供应商-待用原料）\n      Bảng kiểm kê hàng tháng (Bên GC)', '加工方'],
        ['🟢', '委外加工申请单 / Phiếu yêu cầu GC', '生产部 / P.SX'],
    ]
)

add_heading('固资附加单据 / Chứng từ bổ sung – TSCĐ', level=3)
make_table(
    ['级别', '单据 / Chứng từ', '来源 / Nguồn'],
    [
        ['🔴', '设备验收报告 / Biên bản nghiệm thu TB', '使用/技术部'],
        ['🟠', '安装调试记录 / BB lắp đặt, chạy thử', '工程/技术'],
        ['🟠', '固资卡片 / Thẻ TSCĐ', '财务/设备管理'],
        ['🟢', '设备购置申请单 / Phiếu đề xuất mua TB', '需求部门'],
    ]
)

# ── 3.4 单据流转时序 ──
add_heading('3.4 单据流转时序', level=2)
add_para('3.4 Trình tự luân chuyển chứng từ', italic=True, size=13, color=(0x44,0x72,0xC4))

add_para('10步流程：① PR(请购) → ② PO(采购订单) → ③ 合同签订 → ④a 进口发货(提单/发票/PL) / ④b 本地发货(送货单) → ⑤a 报关/缴税 / ⑤b 收货验收 → ⑥a 到厂验收 / ⑥b 质检 → ⑦ 三单匹配(PO↔收货↔发票) → ⑧ 财务入账 → ⑨ 付款审批 → ⑩ 付款执行', bold=True)

add_para('Trình tự 10 bước: ① PR (Yêu cầu mua) → ② PO (Đơn hàng) → ③ Ký HĐ → ④a Xuất hàng NK (Invoice/PL/BL) / ④b Xuất hàng nội địa (Phiếu giao) → ⑤a Khai HQ, nộp thuế / ⑤b Nhận hàng, kiểm nhận → ⑥a Nghiệm thu tại xưởng / ⑥b Kiểm chất lượng → ⑦ Đối chiếu 3 chứng từ (PO ↔ Nhận hàng ↔ Hóa đơn) → ⑧ Hạch toán kế toán → ⑨ Phê duyệt thanh toán → ⑩ Thực hiện thanh toán', italic=True, size=10, color=(0x44,0x72,0xC4))

add_heading('节点责任表 / Bảng phân công trách nhiệm', level=3)
make_table(
    ['节点 Bước', '责任部门 BP', '产出 Đầu ra', '传递给 Chuyển cho'],
    [
        ['① 请购 / Yêu cầu mua', '需求部门 / BP yêu cầu', 'PR', '采购部'],
        ['② 下单 / Đặt hàng', '采购部 / P.Mua hàng', 'PO', '供应商+仓库'],
        ['③ 签约 / Ký HĐ', '采购部 / P.Mua hàng', '合同 / HĐ', '财务(关键页)'],
        ['④ 发货 / Giao hàng', '供应商 / NCC', '发货凭证', '仓库'],
        ['⑤ 报关 / Khai HQ', '报关行 / Đại lý HQ', '报关单+缴款书', '财务'],
        ['⑤/⑥ 验收 / Nghiệm thu', '仓库+质检 / Kho+QC', '入库单+质检报告', '财务'],
        ['⑦ 对单 / Đối chiếu', '采购部', '三单匹配表', '财务'],
        ['⑧ 入账 / Hạch toán', '财务部 / P.Kế toán', '会计凭证', '—'],
        ['⑨ 审批 / Phê duyệt', '管理层 / Ban LĐ', '付款批准', '财务'],
        ['⑩ 付款 / Thanh toán', '财务部 / P.Kế toán', '付款凭证', '供应商'],
    ]
)

# ═══════════════════════════════════════════════
# 第四章
# ═══════════════════════════════════════════════
add_heading('第四章 · 期末处理 & 特殊场景', level=1)

add_heading('4.1 暂估入库（月末货到票未到）', level=2)
add_para('原则：控制权转移>发票到达。货到了就是公司的，应付义务已产生。')
add_entry_block(
    ['月末暂估:  借: 152/211 [按PO不含税价]    贷: 338 - Phải trả tạm tính  ← 不含VAT'],
    ['SOP建议方案B（保留待调），但须配合暂估台账：PO号+暂估金额+日期，发票到后核对调整']
)
add_warning('SOP建议选方案B（保留待调），但必须配合暂估台账跟踪差异')

add_heading('4.2 预付核销 / 4.3 外币应付重估 / 4.4 退货与质量折让', level=2)
add_para('详见原版文档')

add_heading('4.5 委外加工专项 / 4.6 固定资产生命周期衔接', level=2)
add_para('详见原版文档')

# ═══════════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════════
add_heading('📊 附录A · 类别×来源对比矩阵', level=1)
make_table(
    ['维度', '原材料本地', '原材料进口', '生物质本地', '化学品本地', '委外本地', '固资本地', '固资进口'],
    [
        ['资产科目', '152', '152', '152', '152', '154', '211', '211'],
        ['进项税', '1331', '1331', '1331', '1331', '1331', '1332', '1332'],
        ['入账时点', '收货验收', '报关放行', '收货过磅', '收货验收', '发票日', '验收合格', '安装完成'],
        ['成本构成', '买价+运费', 'CIF+关税+港杂', '+含水率', '+环保税', '三段式', '+安装调试', '+关税+安装'],
    ]
)

add_heading('📎 附录B · 案例库', level=1)
add_para('案例1-9详见原版文档（FOB进口分摊、CIF关税拆分、暂估差异、委外场景、质量折让、汇率重估等）')

# ── 保存 ──
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, '..', 'reports', '应付采购入账SOP框架-v5.0-中越双语-20260730.docx')
output_path = os.path.abspath(output_path)
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'✅ 已保存: {output_path}')
