#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 越南差旅费管理办法 v1.0.4：融入征集意见稿6条采纳意见"""
import shutil
from docx import Document
from docx.shared import Pt

SRC = 'reports/越南差旅费管理办法-v1.0.3-20260722.docx'
DST = 'reports/越南差旅费管理办法-v1.0.4-20260805.docx'
shutil.copy(SRC, DST)

d = Document(DST)

def set_para_text(p, new_text):
    """保留段落第一个run格式，替换全部文本"""
    if not p.runs:
        p.add_run(new_text)
        return
    first = p.runs[0]
    first.text = new_text
    for r in p.runs[1:]:
        r.text = ''

def find_para(prefix):
    for p in d.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None

# ── 1. 版本信息 ──
p = find_para('文件编号')
set_para_text(p, '文件编号：VN-CW-01  |  版本：V1.0.4  |  日期：2026-08-05  |  状态：征集意见修订版')

# ── 2. 5.2.1 差旅费申请（采纳 赵立峰：OA/邮件回复同意后方可发起报销）──
p = find_para('5.2.1')
set_para_text(p, '5.2.1 员工因公司业务出差须事前提交出差申请，列明目的地、行程日期、出差事由、预计交通方式、住宿需求、是否需要预支差旅费。差旅费报销流程发起前，须OA出差申请流程结束，或邮件出差申请获直属上司回复同意后，方可发起报销。')

# ── 3. 5.4.1 取消超期惩罚（采纳 陈剑锋）──
p = find_para('5.4.1')
set_para_text(p, '5.4.1 出差行程结束后，须在15个工作日内提交差旅费报销。（取消原"超期限报销按逾期时间延迟付款"惩罚条款）')

# ── 4. 5.4.2 当日出差单据（部分采纳 谢明蓉：仅误餐补贴只需审批记录）──
p = find_para('5.4.2')
set_para_text(p, '5.4.2 报销需提供：出差申请审批记录、交通票据（机票行程单/火车票/网约车行程截图）、住宿发票及其他凭证。如仅为报销误餐补贴，仅需出差审批记录，不需出入打卡证明。')

# ── 5. 5.7.1 市内误餐补贴（采纳 武氏金莲 越南语修正 + 陈剑锋 时间判定）──
p = find_para('5.7.1')
set_para_text(p, '5.7.1 市内（同奈仁泽和西宁展鹏）误餐补贴（₫/天）/ Phụ cấp tiền ăn đối với công tác nội địa phương (Nhơn Trạch, Đồng Nai và Công ty tại Trảng Bàng, Tây Ninh) (₫/ngày)：')
p = find_para('适用条件')
set_para_text(p, '适用条件：7:00前出发→早餐；12:30后未返→午餐；17:30后未返→晚餐。（征集意见修订：原6:30/12:30/18:30 → 7:00/12:30/17:30）')

# ── 6. 权责表 出差员工：超标部分经事业部负责人审批后可报销（采纳 陈剑锋）──
t0 = d.tables[0]
for r in t0.rows:
    cells = [c.text for c in r.cells]
    if '出差员工' in cells[0]:
        c = r.cells[1]
        # 清空后重写
        for para in c.paragraphs:
            for run in para.runs:
                run.text = ''
        c.paragraphs[0].runs[0].text = '提前办理出差申请、按标准预订、一次性报销；超标部分经事业部负责人审批后可报销\nĐăng ký trước, đặt vé/khách sạn theo tiêu chuẩn, thanh toán 1 lần; phần vượt tiêu chuẩn được thanh toán sau khi Trưởng khối duyệt'
        # 删除多余段落
        for para in c.paragraphs[1:]:
            para._element.getparent().remove(para._element)

# ── 7. 附表1 城市等级：二类新增城市（采纳 阮氏云）──
# Table 2: 住宿标准表 二类代表性城市
t2 = d.tables[2]
for r in t2.rows:
    cells = [c.text for c in r.cells]
    if '二类' in cells[0]:
        c = r.cells[1]
        for para in c.paragraphs:
            for run in para.runs:
                run.text = ''
        c.paragraphs[0].runs[0].text = 'Hải Phòng, Huế, Đà Nẵng, Cần Thơ, Thủ Đức, An Giang, Vĩnh Long, Bến Tre, Tiền Giang, Nam Định, Ninh Bình'
        for para in c.paragraphs[1:]:
            para._element.getparent().remove(para._element)
        # 备注行：动态调整机制
        note = d.paragraphs[0]._element
        # 在表格后添加备注段落
        from docx.oxml.ns import qn
        new_p = d.add_paragraph()
        new_p.text = ''
        run = new_p.add_run('📌 征集意见（阮氏云，已采纳）：以上新增城市先按二类分类；后续如发现普遍报销住宿费低于800,000₫，修改为三类城市。')
        run.font.size = Pt(9)

# Table 6: 城市名单表
t6 = d.tables[6]
for r in t6.rows:
    cells = [c.text for c in r.cells]
    if '二类' in cells[0]:
        c = r.cells[1]
        for para in c.paragraphs:
            for run in para.runs:
                run.text = ''
        c.paragraphs[0].runs[0].text = 'Hải Phòng / Huế / Đà Nẵng / Cần Thơ / Thủ Đức / An Giang / Vĩnh Long / Bến Tre / Tiền Giang / Nam Định / Ninh Bình'
        for para in c.paragraphs[1:]:
            para._element.getparent().remove(para._element)
        c2 = r.cells[2]
        for para in c2.paragraphs:
            for run in para.runs:
                run.text = ''
        c2.paragraphs[0].runs[0].text = '海防市 / 顺化市 / 岘港市 / 芹苴市 / 守德郡 / 安江省 / 永隆省 / 槟知省 / 前江省 / 南定省 / 宁平省'
        for para in c2.paragraphs[1:]:
            para._element.getparent().remove(para._element)

d.save(DST)
print('✅ v1.0.4 saved:', DST)
