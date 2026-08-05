#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""v1.0.5 基于 Daryl 8/4 双语版重做：融入征集意见 6 条采纳意见，保留 Daryl 全部翻译与修改"""
import shutil
from docx import Document
from docx.shared import Pt, RGBColor

SRC = '/Users/zhaoyuzhao/.openclaw/media/inbound/越南差旅费管理办法QUY_ĐỊNH_QUẢN_LÝ_CHI_PHÍ_CÔNG_TÁC_TẠI_VIỆT_NAM---6f3dc60e-dc36-4e67-9544-b132677507a8.docx'
DST = 'reports/越南差旅费管理办法-v1.0.5-20260805.docx'
shutil.copy(SRC, DST)
d = Document(DST)

def set_para(p, text):
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ''

def set_cell(cell, text):
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p0 = cell.paragraphs[0]
    if p0.runs:
        p0.runs[0].text = text
        for r in p0.runs[1:]:
            r.text = ''
    else:
        p0.add_run(text)
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)

def find_para(prefix):
    for p in d.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None

# ═══ 0. 版本信息行（标题下方插入）═══
ver_p = d.paragraphs[2].insert_paragraph_before('')
run = ver_p.add_run('版本：V1.0.5  |  日期：2026-08-05  |  状态：征集意见修订版（含8/4双语版全部修改）')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)

# ═══ 1. 4.2.1 赵立峰：OA/邮件回复同意后方可发起报销 ═══
p = find_para('4.2.1')
set_para(p, '''4.2.1 员工因公司业务出差须事前提交出差申请（邮件/OA/书面），列明出差目的地、行程日期、出差事由、预计交通方式、住宿需求，以及是否需要预支差旅费。差旅费报销流程发起前，须OA出差申请流程结束，或邮件出差申请获直属上司回复同意后，方可发起报销。
4.2.1 Nhân viên đi công tác vì hoạt động kinh doanh của công ty phải nộp đề nghị đi công tác trước chuyến đi (qua email/OA/văn bản), trong đó nêu rõ địa điểm, thời gian, lý do đi công tác, phương tiện dự kiến, nhu cầu lưu trú và nhu cầu tạm ứng chi phí công tác. Trước khi lập hồ sơ thanh toán, phải hoàn tất quy trình đề nghị đi công tác trên OA, hoặc được người quản lý trực tiếp phê duyệt qua email.''')

# ═══ 2. 4.4.1 陈剑锋：取消超期延迟付款惩罚 ═══
p = find_para('4.4.1')
set_para(p, '''4.4.1 出差行程结束后，须在15个工作日内提交差旅费报销。
4.4.1 Trong vòng 15 ngày làm việc kể từ khi kết thúc chuyến công tác, nhân viên phải nộp hồ sơ thanh toán chi phí công tác.''')

# ═══ 3. 4.4.2 谢明蓉（部分采纳）：仅误餐补贴只需审批记录 ═══
p = find_para('4.4.2')
set_para(p, '''4.4.2 差旅费报销需提供：出差申请审批记录、交通票据（机票行程单/火车票/网约车行程截图）、住宿发票及其他相关费用凭证。如仅为报销误餐补贴，仅需出差审批记录，不需出入打卡证明。
4.4.2 Hồ sơ thanh toán chi phí công tác phải gồm: hồ sơ phê duyệt đề nghị đi công tác, chứng từ đi lại (xác nhận hành trình vé máy bay/vé tàu hỏa/ảnh chụp hành trình xe công nghệ), hóa đơn lưu trú và các chứng từ chi phí liên quan khác. Nếu chỉ thanh toán phụ cấp tiền ăn, chỉ cần hồ sơ phê duyệt đề nghị đi công tác, không cần xác nhận chấm công ra vào cổng.''')

# ═══ 4. 4.7.1 武氏金莲：越南语修正 Zhanpeng → Trảng Bàng ═══
p = find_para('4.7.1')
set_para(p, '''4.7.1 市内（同奈仁泽和西宁展鹏）误餐补贴（₫/天）：
4.7.1 Phụ cấp tiền ăn đối với công tác nội địa phương (Nhơn Trạch, Đồng Nai và Công ty tại Trảng Bàng, Tây Ninh) (₫/ngày):''')

# ═══ 5. 4.7.1 适用条件 陈剑锋：7:00/12:30/17:30 ═══
p = find_para('适用条件')
set_para(p, '''适用条件：7:00前出发可支付早餐；中午12:30后未返回且未在公司用餐，可支付午餐；晚上17:30后未返回且未在公司用餐，可支付晚餐；
Điều kiện áp dụng: Khởi hành trước 7:00 được thanh toán bữa sáng; sau 12:30 chưa trở về và không dùng bữa tại công ty được thanh toán bữa trưa; sau 17:30 chưa trở về và không dùng bữa tại công ty được thanh toán bữa tối.''')

# ═══ 6. 4.7.4 阈值 10h→12h（与明细表一致）═══
p = find_para('4.7.4')
set_para(p, '''4.7.4 员工出差当天往返，时间超过12小时，可全额报销当天的误餐补贴。
4.7.4 Trường hợp nhân viên đi và về trong ngày, với tổng thời gian trên 12 giờ, được thanh toán toàn bộ phụ cấp tiền ăn trong ngày.''')

# ═══ 7. 权责表 出差员工：超标部分经事业部总经理审批后可报销（陈剑锋）═══
t0 = d.tables[0]
for r in t0.rows:
    if '出差员工' in r.cells[0].text:
        set_cell(r.cells[1], '''提前办理出差申请、按标准预订、一次性报销、超标部分经事业部总经理审批后可报销
Làm thủ tục đề nghị đi công tác trước chuyến đi, đặt dịch vụ đúng tiêu chuẩn, thanh toán một lần; phần vượt tiêu chuẩn được thanh toán sau khi được Tổng Giám đốc Khối phê duyệt''')
        break

# ═══ 8. 附件1（Table 2 住宿标准）二类城市新增（阮氏云）═══
t2 = d.tables[2]
for r in t2.rows:
    if '二类' in r.cells[0].text:
        old = r.cells[1].text
        new = old.rstrip() + ', An Giang安江, Vĩnh Long永隆, Bến Tre槟知, Tiền Giang前江, Nam Định南定, Ninh Bình宁平'
        set_cell(r.cells[1], new)
        break

# ═══ 9. 附件1（Table 6 城市列表）二类城市新增 ═══
t6 = d.tables[6]
for r in t6.rows:
    if '二类' in r.cells[0].text:
        old = r.cells[1].text
        new = old.rstrip() + ', An Giang安江, Vĩnh Long永隆, Bến Tre槟知, Tiền Giang前江, Nam Định南定, Ninh Bình宁平'
        set_cell(r.cells[1], new)
        break

# ═══ 10. 附件1 末尾追加采纳备注（阮氏云机制）═══
last = d.paragraphs[-1]
note = d.add_paragraph('')
rn = note.add_run('📌 征集意见（阮氏云，已采纳）：以上新增城市（安江/永隆/槟知/前江/南定/宁平）先按二类城市分类；后续如发现普遍报销住宿费低于800,000₫，再修改为三类城市。')
rn.font.size = Pt(9)
rn.font.color.rgb = RGBColor(0x99, 0x33, 0x33)

d.save(DST)
print('✅ v1.0.5（Daryl双语版基底）saved:', DST)
