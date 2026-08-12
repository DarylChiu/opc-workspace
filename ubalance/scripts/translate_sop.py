#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
仓储突击检查SOP (Daryl 2026-08-11版) → 中越双语
规则: 中文后换行追加越南语; 不添加目录; 不破坏任何格式/表格结构
"""
import shutil
import docx
from docx.oxml.ns import qn

SRC = '/Users/zhaoyuzhao/.openclaw/media/inbound/仓储突击检查SOP-Daryl-20260811---6fa7d415-440e-4946-b91b-e9c36d59f35b.docx'
OUT = '/Users/zhaoyuzhao/.openclaw/workspace-balance/reports/仓储突击检查SOP-中越双语-20260812.docx'

TRANS = {
    # ===== 标题与正文段落 =====
    '仓储突击检查SOP': 'SOP Kiểm tra Đột xuất Kho hàng',
    '盘点选样指引——选样基础和规则': 'Hướng dẫn chọn mẫu kiểm kê — Cơ sở và quy tắc chọn mẫu',
    '核心逻辑：①抽验样本有效；②尽可能降低抽盘需要的时间；③抽验记录可追踪；':
        'Logic cốt lõi: ①Mẫu kiểm tra có giá trị; ②Giảm tối đa thời gian kiểm kê mẫu; ③Hồ sơ kiểm tra có thể truy vết.',
    'A1.原材料（纱线）、化工仓按环思：有效仓库位置+物料号不少于20个筛选后，按物料号透视表的总数作为目标盘点数量；（见演示）':
        'A1. Kho nguyên vật liệu (sợi) và kho hóa chất theo hệ thống HuanSi: sau khi lọc vị trí kho hợp lệ + mã vật tư ≥ 20, lấy tổng số theo bảng Pivot theo mã vật tư làm số lượng kiểm kê mục tiêu; (xem demo)',
    'A2.坯布和成品仓：有效仓库位置+料号不少于20个筛选后，根据库位前往仓库使用扫描枪扫码库位的所有存货，并在盘点表上记录数量作为扫描完整的依据；（见演示）':
        'A2. Kho vải mộc và kho thành phẩm: sau khi lọc vị trí kho hợp lệ + mã vật tư ≥ 20, căn cứ vị trí kho đến kho dùng máy quét quét toàn bộ hàng tồn tại vị trí đó, ghi số lượng vào bảng kiểm kê làm căn cứ xác nhận đã quét đầy đủ; (xem demo)',
    'A3.备件仓和行政仓：BIP仓库实时数据导出随机抽取不少于20个；':
        'A3. Kho phụ tùng và kho hành chính: xuất dữ liệu thời gian thực từ BIP, chọn ngẫu nhiên ≥ 20 mã;',
    '盘点选样表和排班表（指定组长汇报）由部门主管（可授权）控制与维护，每周突击盘点前企微同步':
        'Bảng chọn mẫu kiểm kê và bảng phân ca (do tổ trưởng được chỉ định báo cáo) do trưởng bộ phận (có thể ủy quyền) kiểm soát và duy trì, đồng bộ qua WeChat Doanh nghiệp (WeCom) trước mỗi đợt kiểm kê đột xuất hàng tuần.',
    'B · 盘点动作指引——5模块问题发现标准动作':
        'B · Hướng dẫn thao tác kiểm kê — Quy trình chuẩn 5 module phát hiện vấn đề',
    '核心逻辑: 不是「去看看有没有问题」，而是每到一个仓库就按以下顺序走一遍。5个模块的顺序是固定的（由外到内、由表到账），形成肌肉记忆。':
        'Logic cốt lõi: Không phải "đi xem có vấn đề gì không", mà mỗi khi đến một kho là đi qua đủ các bước theo thứ tự dưới đây. Thứ tự 5 module là cố định (từ ngoài vào trong, từ hiện trường đến sổ sách), tạo thành phản xạ tự nhiên.',
    'B1. 分区定置标识（到仓第一眼——看分区、看标识、看危险品）':
        'B1. Phân khu và biển định vị (ấn tượng đầu tiên khi đến kho — xem phân khu, xem biển báo, xem hàng nguy hiểm)',
    '检查逻辑: 站在仓库入口→沿通道走一遍→问题应该「扎眼」':
        'Logic kiểm tra: đứng ở cửa kho → đi dọc lối đi một vòng → vấn đề phải "lộ rõ ngay"',
    '常见问题速查: ①柴油桶放五金仓角落无标识 ②油墨/稀释剂无防漏措施 ③货架标签用的是临时手写纸片':
        'Tra nhanh vấn đề thường gặp: ①Thùng dầu diesel để góc kho phụ tùng kim khí không nhãn ②Mực in/dung môi pha loãng không có biện pháp chống rò rỉ ③Nhãn kệ dùng giấy viết tay tạm thời',
    'B2. 现场周转控制（走动中观察——看状态、看日期、看呆滞）':
        'B2. Kiểm soát luân chuyển hiện trường (quan sát khi di chuyển — xem tình trạng, xem ngày tháng, xem hàng tồn chậm)',
    '检查逻辑: 边走边看物料的新旧程度、保质期、堆积时间':
        'Logic kiểm tra: vừa đi vừa xem độ mới/cũ của vật tư, hạn sử dụng, thời gian tồn đọng',
    '常见问题速查: ①化工类过期品OA已报废但实物未处置 ②退货区货物堆放无时间记录 ③呆滞品台账与实际对不上':
        'Tra nhanh vấn đề thường gặp: ①Hàng hóa chất quá hạn đã duyệt hủy trên OA nhưng hiện vật chưa xử lý ②Khu trả hàng chất đống không ghi thời gian ③Sổ theo dõi hàng tồn chậm không khớp thực tế',
    'B3. 收发存控制（到缓冲区/交接区——看交接、看管控、看记录）':
        'B3. Kiểm soát nhập-xuất-tồn (đến khu đệm/khu bàn giao — xem bàn giao, xem kiểm soát, xem ghi chép)',
    '检查逻辑: 聚焦「货从哪来、怎么交接、没交接完的怎么管」':
        'Logic kiểm tra: tập trung vào "hàng từ đâu đến, bàn giao thế nào, hàng chưa bàn giao xong được quản lý ra sao"',
    '常见问题速查: ①缓冲区夜间无管控 ②送货单上只有供应商签字无仓库签字 ③无单领料后长期不补单':
        'Tra nhanh vấn đề thường gặp: ①Khu đệm không kiểm soát vào ban đêm ②Phiếu giao hàng chỉ có chữ ký nhà cung cấp, không có chữ ký kho ③Lĩnh vật tư không phiếu rồi lâu ngày không bổ sung phiếu',
    'B4. 仓库数据维护（坐到保管员电脑前——看系统、看单据、看时效）':
        'B4. Duy trì dữ liệu kho (ngồi vào máy tính thủ kho — xem hệ thống, xem chứng từ, xem tính kịp thời)',
    '检查逻辑: 账实相符的前提是数据录入及时。不看全部数据，只抽最近几天的。':
        'Logic kiểm tra: tiền đề để sổ sách khớp thực tế là nhập liệu kịp thời. Không xem toàn bộ dữ liệu, chỉ lấy mẫu vài ngày gần nhất.',
    '常见问题速查: ①入库频率低的仓库减少录入频次导致系统延迟 ②单据归档混乱（入库单和出库单混放、缺少日期排序）③自盘记录缺失':
        'Tra nhanh vấn đề thường gặp: ①Kho có tần suất nhập thấp giảm tần suất nhập liệu gây chậm trễ hệ thống ②Lưu trữ chứng từ lộn xộn (phiếu nhập và phiếu xuất để lẫn, thiếu sắp xếp theo ngày) ③Thiếu biên bản tự kiểm kê',
    'B5. 盘点管理（回到盘点本身——看配合、看流程、看复盘）':
        'B5. Quản lý kiểm kê (trở lại việc kiểm kê — xem phối hợp, xem quy trình, xem tổng kết rút kinh nghiệm)',
    '检查逻辑: 这部分是「对这次检查本身的复盘」，也是下次改进的依据':
        'Logic kiểm tra: phần này là "tổng kết chính đợt kiểm tra này", cũng là căn cứ để cải tiến lần sau',
    '常见问题速查: ①财务选料号太多导致超时 ②扫码枪没电/故障/不兼容 ③高位货架物料无法清点':
        'Tra nhanh vấn đề thường gặp: ①Bộ phận tài chính chọn quá nhiều mã vật tư gây quá thời gian ②Máy quét hết pin/hỏng/không tương thích ③Vật tư trên kệ cao không đếm được',
    '动作检查顺序总结': 'Tổng kết thứ tự các bước kiểm tra',
    '到仓 → B1 走一圈看标识和危险品（5分钟）\n        → B2 边走边看呆滞过期（5分钟）\n        → B3 到缓冲区看交接管控（5分钟）\n        → B4 坐到电脑前抽单据台账（10分钟）\n        → B5 盘点结束记配合度+复盘（最后5分钟）\n     \n总耗时控制：30分钟/仓（不含物料逐一清点时间）\n物料清点时间：按抽样数量另计':
        'Đến kho → B1 đi một vòng xem biển báo và hàng nguy hiểm (5 phút)\n→ B2 vừa đi vừa xem hàng tồn chậm, quá hạn (5 phút)\n→ B3 đến khu đệm xem bàn giao, kiểm soát (5 phút)\n→ B4 ngồi vào máy tính lấy mẫu chứng từ, sổ sách (10 phút)\n→ B5 sau kiểm kê ghi mức phối hợp + tổng kết (5 phút cuối)\nTổng thời gian: 30 phút/kho (không gồm thời gian đếm từng vật tư)\nThời gian đếm vật tư: tính riêng theo số lượng mẫu',
    'C · 盘点报告汇报与整改进度跟进':
        'C · Báo cáo kiểm kê và theo dõi tiến độ khắc phục',
    '核心原则: 检查完不汇报=没检查。汇报不走闭环=白检查。':
        'Nguyên tắc cốt lõi: Kiểm tra xong không báo cáo = không kiểm tra. Báo cáo không khép kín = kiểm tra vô ích.',
    'C1. 汇报架构': 'C1. Cơ cấu báo cáo',
    '越南财务部(执行层)                               中国总部(监督层)\n     │                                                             │\n     ├─ 检查组长 ──→ 邮件报告 ──→仓储部负责人， 事业部财务负责人\n     │                                     +                                     + 抄送: 副总裁：林金建\n     │                               整改清单                             + 升级抄送: 事业部总经理\n     │\n     ├─ 检查组长 ──→ 口头简报 ──→ 仓库负责人（当场）\n     │\n     └─ 检查组长 ──→ 整改跟踪表 ──→ 共享文档（持续更新）':
        'Bộ phận Tài chính VN (thực thi) → Trụ sở Trung Quốc (giám sát)\nTổ trưởng kiểm tra → Email báo cáo → Trưởng bộ phận kho + Trưởng tài chính đơn vị kinh doanh; CC: Phó Tổng Giám đốc Lâm Kim Kiến; CC nâng cấp: Tổng Giám đốc đơn vị kinh doanh\nTổ trưởng kiểm tra → Báo cáo miệng → Người phụ trách kho (tại chỗ)\nTổ trưởng kiểm tra → Bảng theo dõi khắc phục → Tài liệu dùng chung (cập nhật liên tục)',
    'C2. 汇报节奏': 'C2. Nhịp độ báo cáo',
    'C3. 邮件报告模板': 'C3. Mẫu báo cáo qua email',
    '汇报结构：事由与结果→问题和亮点→未决+证据':
        'Cấu trúc báo cáo: Sự việc và kết quả → Vấn đề và điểm sáng → Tồn đọng + Bằng chứng',
    '主题：【仓储突击检查】[仓库名称] 第[ ]期 [YYYY-MM-DD]\n\n收件人：仓储部负责人、事业部财务负责人\n抄送：林金建、[事业部负责人]、黄云卿、林华军、盘点组员、会计长、仓储课长\n\n正文：\n\n各位领导同事：\n\n一、事由与执行':
        'Chủ đề: 【Kiểm tra đột xuất kho】[Tên kho] Kỳ thứ [ ] [YYYY-MM-DD]\n\nNgười nhận: Trưởng bộ phận kho, Trưởng tài chính đơn vị kinh doanh\nCC: Lâm Kim Kiến, [Người phụ trách đơn vị kinh doanh], Hoàng Vân Khanh, Lâm Hoa Quân, thành viên tổ kiểm kê, Kế toán trưởng, Trưởng ban kho\n\nNội dung:\n\nKính gửi các lãnh đạo, đồng nghiệp:\n\nI. Sự việc và thực hiện',
    '财务部(本次盘点组组长)[姓名]，\n于[YYYY年MM月DD日对[仓库名称]进行突击盘点。\n本次抽查[X]个料号，账实相符[X]项，差异[X]项。\n\n二、发现问题\n按5模块检查结果汇总：范例':
        'Phòng Tài chính (Tổ trưởng tổ kiểm kê đợt này) [Họ tên],\ntiến hành kiểm kê đột xuất kho [Tên kho] vào ngày [YYYY/MM/DD].\nĐợt này kiểm tra mẫu [X] mã vật tư, khớp sổ sách - thực tế [X] mục, chênh lệch [X] mục.\n\nII. Vấn đề phát hiện\nTổng hợp kết quả kiểm tra theo 5 module: ví dụ',
    '\n具体问题与整改要求详见附件《整改事项清单》。\n\n三、未完事项（如有）\n[如有本次未能完成检查的项目，列明原因]\n\n四、附件\n附件1：专项检查盘点表\n附件2：整改事项清单\n附件3：检查照片（问题+亮点）\n\n---\n\n[检查组长姓名]\n[日期]':
        'Chi tiết vấn đề và yêu cầu khắc phục xem Phụ lục 《Danh sách vấn đề cần khắc phục》.\n\nIII. Việc chưa hoàn thành (nếu có)\n[Nếu có hạng mục chưa hoàn thành kiểm tra trong đợt này, nêu rõ lý do]\n\nIV. Phụ lục\nPhụ lục 1: Bảng kiểm kê chuyên đề\nPhụ lục 2: Danh sách vấn đề cần khắc phục\nPhụ lục 3: Ảnh chụp kiểm tra (vấn đề + điểm sáng)\n\n---\n\n[Tên tổ trưởng kiểm tra]\n[Ngày]',
    '具体问题与整改要求详见附件《整改事项清单》。\n\n三、未完事项（如有）\n[如有本次未能完成检查的项目，列明原因]\n\n四、附件\n附件1：专项检查盘点表\n附件2：整改事项清单\n附件3：检查照片（问题+亮点）\n\n---\n\n[检查组长姓名]\n[日期]':
        'Chi tiết vấn đề và yêu cầu khắc phục xem Phụ lục 《Danh sách vấn đề cần khắc phục》.\n\nIII. Việc chưa hoàn thành (nếu có)\n[Nếu có hạng mục chưa hoàn thành kiểm tra trong đợt này, nêu rõ lý do]\n\nIV. Phụ lục\nPhụ lục 1: Bảng kiểm kê chuyên đề\nPhụ lục 2: Danh sách vấn đề cần khắc phục\nPhụ lục 3: Ảnh chụp kiểm tra (vấn đề + điểm sáng)\n\n---\n\n[Tên tổ trưởng kiểm tra]\n[Ngày]',
    'C4. 整改跟踪表（共享文档持续维护）':
        'C4. Bảng theo dõi khắc phục (tài liệu dùng chung, duy trì liên tục)',
    '建议用企微多维表格或共享Excel，字段如下：':
        'Đề xuất dùng bảng đa chiều WeChat Doanh nghiệp hoặc Excel dùng chung, các trường như sau:',

    # ===== 表格通用表头 =====
    '动作编号': 'Mã HĐ',
    '检查动作': 'Hành động kiểm tra',
    '怎么才算有问题': 'Thế nào là có vấn đề',
    '拍照要求': 'Yêu cầu chụp ảnh',
    '节点': 'Mốc thời gian',
    '动作': 'Hành động',
    '时限': 'Thời hạn',
    '产出方式': 'Hình thức đầu ra',
    '序号': 'STT',
    '模块': 'Module',
    '问题&整改建议': 'Vấn đề & kiến nghị khắc phục',
    '严重程度': 'Mức độ nghiêm trọng',
    '字段': 'Trường dữ liệu',
    '填写人': 'Người điền',
    '填写时机': 'Thời điểm điền',
    '说明': 'Ghi chú',

    # ===== B1 表 =====
    '看仓库入口有无平面分区图': 'Xem cửa kho có sơ đồ phân khu mặt bằng không',
    '无分区图，或有图但与实际货位不对应': 'Không có sơ đồ, hoặc có nhưng không khớp vị trí kệ thực tế',
    '入口位置全景': 'Toàn cảnh vị trí cửa vào',
    '沿通道走，看货架/区域有无定置标识牌（如A区-五金、B区-油墨）': 'Đi dọc lối đi, xem kệ/khu vực có biển định vị không (ví dụ Khu A - phụ tùng kim khí, Khu B - mực in)',
    '标识缺失、模糊、掉落未补': 'Biển thiếu, mờ, rơi chưa gắn lại',
    '拍到无标识的区域全景+货架编号（如有）': 'Toàn cảnh khu vực thiếu biển + mã kệ (nếu có)',
    '随机抽查3-5个货位，看物料卡/标签是否在位且信息完整（物料编码+名称+规格+单位）': 'Kiểm tra ngẫu nhiên 3-5 vị trí kệ, xem thẻ/nhãn vật tư có đúng vị trí và đủ thông tin (mã vật tư + tên + quy cách + đơn vị)',
    '无标签、标签信息不全、手写涂改未签字、标签与实物不符': 'Không nhãn, nhãn thiếu thông tin, sửa tay không ký, nhãn không khớp hiện vật',
    '问题标签特写+所在货架全景': 'Cận cảnh nhãn lỗi + toàn cảnh kệ chứa',
    '关键项：扫一遍仓库，看有无危险品/易燃品（柴油、溶剂、化学品）与其他物料混放': 'Hạng mục then chốt: quét một lượt kho, xem có hàng nguy hiểm/dễ cháy (dầu diesel, dung môi, hóa chất) để lẫn với vật tư khác không',
    '①危化品无独立区域 ②无警示标识 ③无消防器材/消防通道被堵 ④化学品无二次防漏托盘': '①Hóa chất nguy hiểm không có khu vực riêng ②Không có biển cảnh báo ③Không có thiết bị PCCC / lối thoát hiểm bị chặn ④Hóa chất không có khay chống rò rỉ thứ cấp',
    '危化品存放位置全景+警示标识缺失特写': 'Toàn cảnh nơi lưu hóa chất nguy hiểm + cận cảnh thiếu biển cảnh báo',
    '看通道是否畅通（主通道≥1.2m，次通道≥0.8m为参考线）': 'Xem lối đi có thông thoáng không (lối chính ≥1.2m, lối phụ ≥0.8m là chuẩn tham chiếu)',
    '货物堆到通道、通道被叉车/托盘长时间占用': 'Hàng hóa chất ra lối đi, lối đi bị xe nâng/pallet chiếm trong thời gian dài',
    '通道堵塞全景': 'Toàn cảnh lối đi bị tắc',

    # ===== B2 表 =====
    '看物料外观：有无明显过期标识（生产日期+保质期已过）、锈蚀/结块/变色/包装破损': 'Xem ngoại quan vật tư: có dấu hiệu quá hạn rõ ràng (ngày sản xuất + hạn dùng đã qua), rỉ sét / vón cục / đổi màu / vỡ bao bì',
    '目视即可判断的过期或变质': 'Quá hạn hoặc hư hỏng có thể nhận biết bằng mắt thường',
    '过期品标签特写+整堆过期品全景': 'Cận cảnh nhãn hàng quá hạn + toàn cảnh đống hàng quá hạn',
    '问保管员：呆滞品清单在哪里？多久更新一次？': 'Hỏi thủ kho: danh sách hàng tồn chậm ở đâu? Bao lâu cập nhật một lần?',
    '①无呆滞品清单 ②有清单但超过3个月未更新 ③清单上的物料现场找不到或反过来': '①Không có danh sách ②Có nhưng hơn 3 tháng chưa cập nhật ③Vật tư trong danh sách không tìm thấy tại hiện trường, hoặc ngược lại',
    '呆滞品区域全景': 'Toàn cảnh khu vực hàng tồn chậm',
    '抽查3-5项呆滞品，看OA报废流程是否已完成': 'Kiểm tra mẫu 3-5 món hàng tồn chậm, xem quy trình hủy trên OA đã hoàn tất chưa',
    'OA已批报废但实物仍在仓（说明处置脱节）': 'OA đã duyệt hủy nhưng hiện vật vẫn trong kho (xử lý chưa khép kín)',
    'OA截图（保管员提供）+实物照片': 'Ảnh chụp màn hình OA (thủ kho cung cấp) + ảnh hiện vật',
    '看退货区/待处理区的物料，问这些放了多久': 'Xem vật tư ở khu trả hàng / khu chờ xử lý, hỏi đã để bao lâu',
    '超过1个月未处理的退货/待处理品且无跟进记录': 'Hàng trả / hàng chờ xử lý quá 1 tháng chưa xử lý và không có hồ sơ theo dõi',
    '退货区全景+最近一次处理记录': 'Toàn cảnh khu trả hàng + biên bản xử lý gần nhất',
    '问保管员近期入库/出库频率，到高频出入区域看地面有无散落、包装残损': 'Hỏi thủ kho tần suất nhập/xuất gần đây, đến khu vực nhập xuất tần suất cao xem mặt sàn có vật tư rơi vãi, bao bì hư hỏng không',
    '地面有散落物料、破损包装未清理': 'Sàn có vật tư rơi vãi, bao bì hỏng chưa dọn dẹp',
    '散落区域全景': 'Toàn cảnh khu vực rơi vãi',

    # ===== B3 表 =====
    '找到入库暂存区/缓冲区，看暂存货物有无临时标识（到货日期+供应商+物料名称）': 'Tìm khu tạm chứa / khu đệm khi nhập kho, xem hàng tạm chứa có nhãn tạm không (ngày đến hàng + nhà cung cấp + tên vật tư)',
    '暂存货物无任何标识，无法区分批次和归属': 'Hàng tạm chứa không có nhãn, không phân biệt được lô hàng và chủ sở hữu',
    '无标识暂存货全景': 'Toàn cảnh hàng tạm chứa không nhãn',
    '问保管员：当日未入库的暂存货怎么管控？下班后谁负责？': 'Hỏi thủ kho: hàng tạm chứa chưa nhập kho trong ngày được quản lý thế nào? Sau giờ làm ai phụ trách?',
    '①无管控措施（无人锁门/无交接记录） ②下班后缓冲区可自由进出': '①Không có biện pháp kiểm soát (không khóa cửa / không biên bản bàn giao) ②Sau giờ làm khu đệm ra vào tự do',
    '缓冲区与外部通道连接处全景': 'Toàn cảnh điểm nối giữa khu đệm và lối đi bên ngoài',
    '看收发交接记录（送货单签收联/入库单），抽查最近3天': 'Xem biên bản bàn giao nhận/xuất (liên ký nhận phiếu giao hàng / phiếu nhập kho), lấy mẫu 3 ngày gần nhất',
    '①无交接记录 ②交接记录无双方签字 ③到货日期与入库日期超过2个工作日且无说明': '①Không có biên bản bàn giao ②Biên bản không đủ chữ ký hai bên ③Ngày đến hàng và ngày nhập kho cách hơn 2 ngày làm việc mà không có giải trình',
    '缺失的记录台账': 'Sổ ghi chép bị thiếu hụt',
    '看出库是否有领料单，抽查最近3天的领料单': 'Xem xuất kho có phiếu lĩnh vật tư không, lấy mẫu phiếu lĩnh 3 ngày gần nhất',
    '①无领料单 ②领料单无审批签字 ③领用数量与系统扣减数量不一致': '①Không có phiếu lĩnh ②Phiếu lĩnh không có chữ ký phê duyệt ③Số lượng lĩnh không khớp số lượng hệ thống đã trừ',
    '问题单据特写': 'Cận cảnh chứng từ lỗi',
    '问保管员：紧急领料（无单先领）怎么处理？补单时限？': 'Hỏi thủ kho: lĩnh vật tư khẩn cấp (lĩnh trước không phiếu) xử lý thế nào? Thời hạn bổ sung phiếu?',
    '①有制度未执行（抽查到无单领料且超期未补）': '①Có quy định nhưng không thực hiện (lấy mẫu phát hiện lĩnh không phiếu và quá hạn chưa bổ sung)',
    '相关记录': 'Biên bản liên quan',

    # ===== B4 表 =====
    '打开仓库入库单台账，抽查最近一周': 'Mở sổ theo dõi phiếu nhập kho, lấy mẫu tuần gần nhất',
    '①无台账 ②台账不是连续的（有跳跃编号）③入库单与台账日期不一致': '①Không có sổ ②Sổ không liên tục (số bị nhảy) ③Phiếu nhập và sổ lệch ngày',
    '台账翻阅全景': 'Toàn cảnh lật sổ theo dõi',
    '打开仓库出库/领料单台账，抽查最近一周': 'Mở sổ theo dõi phiếu xuất / phiếu lĩnh kho, lấy mẫu tuần gần nhất',
    '①无台账 ②领料单未在次日12:00前录入系统（或已录入但无签名确认）': '①Không có sổ ②Phiếu lĩnh chưa nhập hệ thống trước 12:00 ngày hôm sau (hoặc đã nhập nhưng chưa ký xác nhận)',
    '抽查到的未录入单据特写': 'Cận cảnh chứng từ chưa nhập hệ thống (được lấy mẫu)',
    '抽查3-5张最近入库/出库单，反向追踪到系统（BIP/ERP），看系统数据与单据是否一致': 'Lấy mẫu 3-5 phiếu nhập/xuất gần nhất, truy ngược lên hệ thống (BIP/ERP), xem dữ liệu hệ thống có khớp chứng từ không',
    '单据有、系统无；或系统有、单据找不到': 'Có chứng từ nhưng hệ thống không có; hoặc hệ thống có nhưng không tìm thấy chứng từ',
    '单据vs系统截图对比': 'Ảnh chụp màn hình so sánh chứng từ vs hệ thống',
    '问保管员：自盘记录在哪里？最近一次什么时候？': 'Hỏi thủ kho: biên bản tự kiểm kê ở đâu? Lần gần nhất khi nào?',
    '①无自盘记录 ②最近一次超过1个月 ③自盘差异未跟进处理': '①Không có biên bản tự kiểm kê ②Lần gần nhất hơn 1 tháng ③Chênh lệch tự kiểm kê chưa được theo dõi xử lý',
    '自盘记录本翻阅': 'Lật sổ tự kiểm kê',
    '问保管员：重要内控报告（如月末库存报表、差异分析）归档在哪里': 'Hỏi thủ kho: các báo cáo kiểm soát nội bộ quan trọng (như báo cáo tồn kho cuối tháng, phân tích chênh lệch) được lưu trữ ở đâu',
    '①无归档 ②找不到 ③报告有但缺少当月/当季的': '①Không lưu trữ ②Không tìm thấy ③Có báo cáo nhưng thiếu báo cáo tháng/quý hiện tại',
    '归档位置': 'Vị trí lưu trữ',

    # ===== B5 表 =====
    '记录本次抽盘实际耗时': 'Ghi thời gian thực tế của đợt kiểm kê mẫu này',
    '超过1.5小时→下次选料号要减少；不足0.5小时→下次可增加': 'Quá 1.5 giờ → lần sau giảm số mã vật tư; dưới 0.5 giờ → lần sau có thể tăng',
    '记录仓储部配合情况：是否指定了带路人？扫码枪/叉车/登高设备是否可用？物料是否可触及？': 'Ghi mức phối hợp của bộ phận kho: có cử người dẫn đường không? Máy quét / xe nâng / thiết bị leo cao có sẵn sàng không? Vật tư có lấy tới được không?',
    '①无人带路 ②设备故障影响盘点 ③高位物料无法安全取用': '①Không có người dẫn đường ②Thiết bị hỏng ảnh hưởng kiểm kê ③Vật tư trên cao không thể lấy an toàn',
    '设备故障/无法盘点位置': 'Vị trí thiết bị hỏng / không kiểm kê được',
    '记录盘点中断原因（如有）': 'Ghi nguyên nhân gián đoạn kiểm kê (nếu có)',
    '任何导致抽盘未能100%完成的原因': 'Mọi nguyên nhân khiến kiểm kê mẫu không hoàn thành 100%',
    '盘点结束后与保管员当面确认差异，双方签字': 'Sau khi kiểm kê, xác nhận chênh lệch trực tiếp với thủ kho, hai bên ký tên',
    '保管员拒绝签字→记入报告「配合异常」': 'Thủ kho từ chối ký → ghi vào báo cáo "phối hợp bất thường"',
    '签字后的盘点表': 'Bảng kiểm kê sau khi ký',
    '盘点结束后组内快速复盘（≤5分钟），确认本次检查的3个最值得记录的点（问题/亮点/改进建议各1）': 'Sau kiểm kê, tổ chức tổng kết nhanh trong tổ (≤5 phút), xác nhận 3 điểm đáng ghi nhận nhất của đợt này (vấn đề / điểm sáng / kiến nghị cải tiến, mỗi loại 1)',

    # ===== C2 汇报节奏表 =====
    '当场': 'Tại chỗ',
    '口头简报仓库负责人：本次发现几个问题、严重程度': 'Báo cáo miệng cho người phụ trách kho: đợt này phát hiện bao nhiêu vấn đề, mức độ nghiêm trọng',
    '盘点结束前': 'Trước khi kết thúc kiểm kê',
    '口头': 'Miệng',
    '24h内': 'Trong 24h',
    '组员录入盘点数据到共享文档': 'Thành viên tổ nhập dữ liệu kiểm kê vào tài liệu dùng chung',
    '次日下班前': 'Trước khi tan ca ngày hôm sau',
    '共享文档': 'Tài liệu dùng chung',
    '48h内': 'Trong 48h',
    '组长发邮件报告+整改清单': 'Tổ trưởng gửi email báo cáo + danh sách khắc phục',
    '第2个工作日17:00前': 'Trước 17:00 ngày làm việc thứ 2',
    '邮件+附件': 'Email + phụ lục',
    '持续': 'Liên tục',
    '跟进整改状态，更新跟踪表': 'Theo dõi trạng thái khắc phục, cập nhật bảng theo dõi',
    '实时': 'Thời gian thực',
    '下期检查前': 'Trước đợt kiểm tra tiếp theo',
    '对照上期清单逐项复核': 'Đối chiếu từng mục trong danh sách kỳ trước',
    '检查前1天': '1 ngày trước khi kiểm tra',
    '复核记录': 'Biên bản rà soát lại',

    # ===== C4/整改跟踪表 =====
    '分区定置标识': 'Phân khu và biển định vị',
    '现场周转控制': 'Kiểm soát luân chuyển hiện trường',
    '收发存控制': 'Kiểm soát nhập-xuất-tồn',
    '仓库数据维护': 'Duy trì dữ liệu kho',
    '盘点管理': 'Quản lý kiểm kê',
    '检查日期': 'Ngày kiểm tra',
    '组长': 'Tổ trưởng',
    '检查当天': 'Ngày kiểm tra',
    '检查期数': 'Kỳ kiểm tra thứ mấy',
    '该仓库第X期': 'Kỳ thứ X của kho này',
    '仓库名称': 'Tên kho',
    '保管人': 'Thủ kho',
    '问题编号': 'Mã vấn đề',
    '报告发出时': 'Khi gửi báo cáo',
    '格式：仓库-期数-序号': 'Định dạng: Kho - Kỳ - STT',
    '所属模块': 'Module thuộc về',
    '5模块之一': 'Một trong 5 module',
    '严重级别': 'Mức độ nghiêm trọng',
    '问题描述': 'Mô tả vấn đề',
    '≤30字': '≤30 chữ',
    '整改措施': 'Biện pháp khắc phục',
    '仓库负责人': 'Người phụ trách kho',
    '收到通知后48h内': 'Trong 48h sau khi nhận thông báo',
    '责任人': 'Người chịu trách nhiệm',
    '同上': 'Như trên',
    '承诺完成日': 'Ngày cam kết hoàn thành',
    '实际完成日': 'Ngày hoàn thành thực tế',
    '完成后填写': 'Điền sau khi hoàn thành',
    '整改证据': 'Bằng chứng khắc phục',
    '完成后上传': 'Tải lên sau khi hoàn thành',
    '照片/截图链接': 'Link ảnh / ảnh chụp màn hình',
    '财务复核': 'Tài chính rà soát lại',
    '组长/事业部财务': 'Tổ trưởng / Tài chính đơn vị kinh doanh',
    '复核后': 'Sau khi rà soát',
    '通过/不通过/部分通过': 'Đạt / Không đạt / Đạt một phần',
    '升级标记': 'Đánh dấu nâng cấp',
    '超期未完成时': 'Khi quá hạn chưa hoàn thành',
    '已升级至[XX]': 'Đã nâng cấp lên [XX]',
}

SKIP = set('— 1 2 3 4 5 🔴🟠🟡 YYYY-MM-DD [🔴/🟠/🟡] |'.split()) | {
    '[🔴/🟠/🟡] |',
    'B1-1','B1-2','B1-3','B1-4','B1-5','B2-1','B2-2','B2-3','B2-4','B2-5',
    'B3-1','B3-2','B3-3','B3-4','B3-5','B4-1','B4-2','B4-3','B4-4','B4-5',
    'B5-1','B5-2','B5-3','B5-4','B5-5',
}


def add_vn(paragraph, vn_text):
    """中文后换行追加越南语; 每行越南语前都加 break(用 add_text 保留换行, 不能用 run.text setter)"""
    lines = vn_text.split('\n')
    run = paragraph.add_run()
    for line in lines:
        run.add_break()
        run.add_text(line)


def tc_text(tc):
    """直接读取 tc 元素的文本(合并单元格在 XML 中只出现一次, 无重复)"""
    texts = []
    for p in tc.findall(qn('w:p')):
        texts.append(''.join(t.text or '' for t in p.iter(qn('w:t'))))
    return '\n'.join(texts)


def main():
    shutil.copy(SRC, OUT)
    doc = docx.Document(OUT)
    matched, skipped, missing = 0, 0, []

    # 正文段落
    for p in doc.paragraphs:
        key = p.text.strip()
        if not key:
            continue
        if key in TRANS:
            add_vn(p, TRANS[key])
            matched += 1
        elif key in SKIP:
            skipped += 1
        else:
            missing.append(('P', key[:80]))

    # 表格单元格: 直接遍历 w:tc 元素(每个单元格恰好一次, 无需 id 去重)
    from docx.text.paragraph import Paragraph
    for ti, t in enumerate(doc.tables):
        for tc in t._tbl.iter(qn('w:tc')):
            key = tc_text(tc).strip()
            if not key:
                continue
            if key in TRANS:
                first_p = tc.find(qn('w:p'))
                add_vn(Paragraph(first_p, None), TRANS[key])
                matched += 1
            elif key in SKIP:
                skipped += 1
            else:
                missing.append((f'T{ti}', key[:80]))

    doc.save(OUT)
    print(f'✅ matched={matched} skipped={skipped} missing={len(missing)}')
    for m in missing:
        print('  MISSING:', m)
    return len(missing)


if __name__ == '__main__':
    main()
