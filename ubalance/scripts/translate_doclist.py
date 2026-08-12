#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
应付采购单据清单 (Daryl 0810) → 中越双语
规则: 中文后换行追加越南语; 不加目录; 不破坏表格格式
"""
import shutil
import docx
from docx.oxml.ns import qn

SRC = '/Users/zhaoyuzhao/.openclaw/media/inbound/应付采购单据清单-Daryl0810---c7129ee2-9b2c-4903-bf1b-d60207ee69b7.docx'
OUT = '/Users/zhaoyuzhao/.openclaw/workspace-balance/reports/应付采购单据清单-中越双语-20260812.docx'

TRANS = {
    # ===== 段落 =====
    '应付采购单据清单': 'Danh sách chứng từ mua hàng phải trả',
    '3.3 单据完整性清单*': '3.3 Danh sách kiểm tra tính đầy đủ chứng từ*',
    '分级标准': 'Tiêu chuẩn phân cấp',
    '本地采购单据': 'Chứng từ mua hàng trong nước',
    '进口采购单据（包括越南EPE，越南本地保税企业）': 'Chứng từ mua hàng nhập khẩu (bao gồm EPE Việt Nam, doanh nghiệp bảo thuế nội địa Việt Nam)',
    '委外加工单据': 'Chứng từ gia công bên ngoài',
    '化学品采购单据': 'Chứng từ mua hóa chất',
    '生物质材料采购单据': 'Chứng từ mua nguyên liệu sinh khối',
    '固资采购单据': 'Chứng từ mua tài sản cố định',
    '附录：单据流转时序': 'Phụ lục: Trình tự luân chuyển chứng từ',
    '部门': 'Bộ phận',
    # 流转时序图(ASCII艺术保留中文, 追加越南语线性摘要)
    '① PR(请购) ──→ ②合同签订 ──→     ③PO(采购订单)\n                                        │\n                          ┌─────────────┤\n                          ▼             ▼\n                    ④ 进口发货      ④ 本地发货\n                    (提单/发票/PL)   (送货单)\n                          │             │\n                          ▼             ▼\n                    ⑤ （封签）报关/缴税  ⑤ （过磅单）收货验收\n                    (报关单/缴款书)  (入库单)\n                          │             │\n                          ▼             ▼\n                    ⑥ 到厂验收     ⑥ 质检\n                          │             │\n                          └──────┬──────┘\n                                 ▼\n                          ⑦ 三单匹配\n                        (PO ↔ 收货 ↔ 发票)\n                                 │\n                                 ▼\n                          ⑧ 财务入账\n                                 │\n                                 ▼\n                          ⑨ 付款审批\n                                 │\n                                 ▼\n                          ⑩ 付款执行':
        '① PR (Yêu cầu mua) → ② Ký hợp đồng → ③ PO (Đơn đặt hàng)\n④ Nhập khẩu: Giao hàng (Vận đơn / Hóa đơn / PL) → Khai báo hải quan, niêm phong, nộp thuế (Tờ khai / Giấy nộp thuế) → Nghiệm thu tại nhà máy\n④ Nội địa: Giao hàng (Phiếu giao hàng) → Nhận hàng, cân (Phiếu cân) → Nhập kho (Phiếu nhập kho) → Kiểm tra chất lượng\n⑦ Đối chiếu 3 chứng từ (PO ↔ Nhận hàng ↔ Hóa đơn) → ⑧ Hạch toán kế toán → ⑨ Phê duyệt thanh toán → ⑩ Thực hiện thanh toán',

    # ===== 分级标准 =====
    '级别': 'Cấp độ',
    '含义': 'Ý nghĩa',
    '规则': 'Quy tắc',
    '强制性': 'Bắt buộc',
    '缺少不可入账，禁止付款。必须补齐后才能进入下一流程节点。': 'Thiếu không được hạch toán, cấm thanh toán. Phải bổ sung đầy đủ mới được chuyển sang bước tiếp theo của quy trình.',
    '条件性': 'Có điều kiện',
    '缺了可以先入账，但限期补（通常收货后7天内）。超期未补→升级为🔴，暂停后续付款。': 'Thiếu vẫn hạch toán được trước, nhưng phải bổ sung trong thời hạn (thường trong 7 ngày sau khi nhận hàng). Quá hạn chưa bổ sung → nâng lên mức 🔴, tạm dừng thanh toán tiếp theo.',
    '参考性': 'Tham khảo',
    '内控留痕用。缺了不卡流程，系统留痕关联即可，不打印；。': 'Dùng để lưu vết nội kiểm. Thiếu không chặn quy trình, chỉ cần hệ thống lưu vết liên kết, không cần in.',

    # ===== 通用表头 =====
    '单据': 'Chứng từ',
    '来源': 'Nguồn gốc',
    '缺失后果': 'Hậu quả nếu thiếu',

    # ===== 本地采购 =====
    '采购订单(PO)': 'Đơn đặt hàng (PO)',
    '采购部': 'Phòng Mua hàng',
    '无法证明采购经过授权': 'Không chứng minh được việc mua hàng đã được phê duyệt',
    '供应商送货单': 'Phiếu giao hàng của nhà cung cấp',
    '仓库': 'Kho',
    '收发差异无法追溯': 'Không truy vết được chênh lệch nhập/xuất',
    '过磅单': 'Phiếu cân',
    '地磅': 'Trạm cân',
    '无法佐证入库重量/数量': 'Không chứng minh được trọng lượng/số lượng nhập kho',
    '入库单/收货单': 'Phiếu nhập kho / Phiếu nhận hàng',
    '无法确认资产已实际入库': 'Không xác nhận được tài sản đã thực nhập kho',
    'VAT发票(hóa đơn GTGT)': 'Hóa đơn GTGT',
    '供应商': 'Nhà cung cấp',
    '无法抵扣进项税': 'Không khấu trừ được thuế GTGT đầu vào',
    '采购合同(第一次请款打印，后续OA流程关联不打印)': 'Hợp đồng mua hàng (lần đầu yêu cầu thanh toán in ra, các lần sau liên kết trên luồng OA, không in)',
    '付款条件/价格条款无依据': 'Không có căn cứ về điều kiện thanh toán / điều khoản giá',
    '质检报告(按集团规定)': 'Biên bản kiểm tra chất lượng (theo quy định của Tập đoàn)',
    '质检部': 'Phòng QC',
    '质量无保证': 'Chất lượng không được đảm bảo',
    '请购单(PR)': 'Phiếu yêu cầu mua (PR)',
    '系统留痕': 'Hệ thống lưu vết',
    '审批留痕': 'Lưu vết phê duyệt',
    '比价记录(采购比价系统)': 'Biên bản so sánh giá (hệ thống so sánh giá mua)',
    '内控审计': 'Kiểm toán nội bộ',

    # ===== 进口采购 =====
    '采购订单(PO)/Sales Comfirmation': 'Đơn đặt hàng (PO) / Xác nhận bán hàng (Sales Confirmation)',
    '封签/过磅单': 'Phiếu niêm phong / Phiếu cân',
    '仓库检查封签，不打印；过磅单不检查': 'Kho kiểm tra niêm phong, không in; phiếu cân không kiểm tra',
    '形式或商业发票': 'Hóa đơn chiếu lệ hoặc hóa đơn thương mại',
    '物流部': 'Phòng Logistics',
    '装箱单': 'Phiếu đóng gói (Packing List)',
    '提单': 'Vận đơn (Bill of Lading)',
    '报关单': 'Tờ khai hải quan',
    '报关委托书-报关费结算': 'Giấy ủy quyền khai báo hải quan - quyết toán phí khai báo',
    '报关行—不在进口采购单据中': 'Công ty khai hải quan — không thuộc chứng từ mua hàng nhập khẩu',

    # ===== 委外加工 =====
    '委外加工合同': 'Hợp đồng gia công bên ngoài',
    '委外加工申请单': 'Phiếu đề nghị gia công bên ngoài',
    '生产部/规划部': 'Phòng Sản xuất / Phòng Kế hoạch',
    '委外加工订单': 'Đơn hàng gia công bên ngoài',
    '材料出库单(发外)': 'Phiếu xuất kho vật tư (gửi đi gia công)',
    '加工完成入库单': 'Phiếu nhập kho hàng gia công hoàn thành',
    '月度盘点表（每个月委外供应商提供盘点表-待用原料）': 'Bảng kiểm kê tháng (mỗi tháng nhà cung cấp gia công cung cấp bảng kiểm kê - nguyên liệu chờ sử dụng)',
    '加工方': 'Bên gia công',
    '提供采购部统计台账并计算损耗': 'Cung cấp sổ thống kê cho Phòng Mua hàng và tính toán hao hụt',
    '余料退回单': 'Phiếu trả lại nguyên liệu thừa',

    # ===== 化学品 =====
    '过磅单（如为重量计量）': 'Phiếu cân (nếu tính theo trọng lượng)',
    '采购部（检查报告不打印单据，但是需要流程附上）': 'Phòng Mua hàng (báo cáo kiểm tra không in chứng từ, nhưng cần đính kèm trên luồng)',
    '检查报告': 'Báo cáo kiểm tra',
    '集团统一，出现不及格采购扣款': 'Tập đoàn thống nhất, nếu kết quả không đạt sẽ khấu trừ tiền hàng',

    # ===== 生物质 =====
    '检测报告': 'Biên bản kiểm nghiệm',
    '质检部（按合同条款）': 'Phòng QC (theo điều khoản hợp đồng)',
    '不符检测结果扣款计算表': 'Bảng tính khấu trừ khi kết quả kiểm nghiệm không đạt',
    '质检部/采购部': 'Phòng QC / Phòng Mua hàng',

    # ===== 固资 =====
    '设备验收报告': 'Biên bản nghiệm thu thiết bị',
    '使用/技术部': 'Phòng Sử dụng / Kỹ thuật',
    '安装调试记录（工厂支付）': 'Biên bản lắp đặt, chạy thử (nhà máy thanh toán)',
    '工程/技术': 'Kỹ thuật / Công trình',
    '设备购置申请单': 'Phiếu đề nghị mua sắm thiết bị',
    '需求部门': 'Bộ phận có nhu cầu',

    # ===== 流转时序表 =====
    '节点': 'Bước',
    '责任部门': 'Bộ phận chịu trách nhiệm',
    '产出': 'Đầu ra',
    '传递给': 'Chuyển cho',
    '①请购': '① Yêu cầu mua',
    '②下单': '② Đặt hàng',
    '供应商+仓库': 'Nhà cung cấp + Kho',
    '③签约': '③ Ký hợp đồng',
    '合同': 'Hợp đồng',
    '财务(关键页)': 'Kế toán (trang chính)',
    '④发货': '④ Giao hàng',
    '发货凭证': 'Chứng từ giao hàng',
    '⑤报关': '⑤ Khai báo hải quan',
    '报关行': 'Công ty khai hải quan',
    '报关单+缴款书': 'Tờ khai hải quan + Giấy nộp thuế',
    '财务': 'Kế toán',
    '⑤/⑥验收': '⑤/⑥ Nghiệm thu',
    '仓库+质检': 'Kho + QC',
    '入库单+质检报告': 'Phiếu nhập kho + Biên bản QC',
    '⑦对单': '⑦ Đối chiếu chứng từ',
    '三单匹配表': 'Bảng đối chiếu 3 chứng từ',
    '⑧入账': '⑧ Hạch toán',
    '财务部': 'Phòng Kế toán',
    '会计凭证': 'Chứng từ kế toán',
    '⑨审批': '⑨ Phê duyệt',
    '管理层': 'Ban lãnh đạo',
    '付款批准': 'Phê duyệt thanh toán',
    '⑩付款': '⑩ Thanh toán',
    '付款凭证': 'Chứng từ thanh toán',
}

SKIP = set('— 🔴 🟠 🟢 PR PO'.split())


def add_vn(paragraph, vn_text):
    lines = vn_text.split('\n')
    run = paragraph.add_run()
    for line in lines:
        run.add_break()
        run.add_text(line)


def tc_text(tc):
    return '\n'.join(''.join(t.text or '' for t in p.iter(qn('w:t'))) for p in tc.findall(qn('w:p')))


def main():
    shutil.copy(SRC, OUT)
    doc = docx.Document(OUT)
    matched, skipped, missing = 0, 0, []

    from docx.text.paragraph import Paragraph
    for p in doc.paragraphs:
        key = p.text.strip()
        if not key:
            continue
        if key in TRANS:
            add_vn(p, TRANS[key])
            matched += 1
        elif key in SKIP:
            skipped += 1
        else:
            missing.append(('P', key[:70]))

    for ti, t in enumerate(doc.tables):
        for tc in t._tbl.iter(qn('w:tc')):
            key = tc_text(tc).strip()
            if not key:
                continue
            if key in TRANS:
                first_p = tc.find(qn('w:p'))
                add_vn(Paragraph(first_p, None), TRANS[key])
                matched += 1
            elif key in SKIP:
                skipped += 1
            else:
                missing.append((f'T{ti}', key[:70]))

    doc.save(OUT)
    print(f'✅ matched={matched} skipped={skipped} missing={len(missing)}')
    for m in missing:
        print('  MISSING:', m)
    return len(missing)


if __name__ == '__main__':
    main()
