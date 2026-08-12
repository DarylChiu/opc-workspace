#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""仓储突击检查·仓储部配合事项清单 v1.0 (中越双语, 给仓储部看的版本)"""
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/Users/zhaoyuzhao/.openclaw/workspace-balance/reports/仓储突击检查-仓储部配合清单-v1.0-中越双语-20260812.docx'

doc = docx.Document()

# 页面边距 (2cm, 给表格更多宽度)
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2); s.right_margin = Cm(2)

# Normal 样式: ascii=Calibri(越南语), eastAsia=微软雅黑(中文)
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color


def add_bilingual(p, zh, vi, size=10.5, bold=False):
    r1 = p.add_run(zh)
    set_font(r1, size, bold)
    r2 = p.add_run()
    r2.add_break()
    r2.add_text(vi)
    set_font(r2, size)


# ===== 标题 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_bilingual(p, '仓储突击检查 · 仓储部配合事项清单', 'Kiểm tra đột xuất kho · Danh sách việc phối hợp của bộ phận kho', size=16, bold=True)

# ===== 定位说明 =====
p = doc.add_paragraph()
add_bilingual(p,
    '说明：日常管理事项（设备维护、单据台账、自盘、呆滞品、暂存货管控等）属仓储部平时职责，突击检查不作另行通知；本清单仅列检查过程中的配合动作与检查后的整改闭环。',
    'Ghi chú: Các công việc quản lý hàng ngày (bảo trì thiết bị, sổ sách chứng từ, tự kiểm kê, hàng tồn chậm, quản lý hàng tạm chứa...) thuộc trách nhiệm thường xuyên của bộ phận kho, kiểm tra đột xuất không thông báo trước; danh sách này chỉ liệt kê các thao tác phối hợp trong quá trình kiểm tra và khép kín khắc phục sau kiểm tra.',
    size=9)

# ===== 表头与内容 =====
headers = [('序号', 'STT'), ('配合事项', 'Nội dung phối hợp'), ('具体要求', 'Yêu cầu cụ thể'), ('责任人', 'Người chịu trách nhiệm'), ('时限', 'Thời hạn')]

rows_data = [
    ('1',
     ('人员陪同', 'Người đi kèm'),
     ('检查时指定熟悉库位的人员全程陪同（带路、取货、答疑），不得拒绝配合',
      'Chỉ định nhân viên am hiểu vị trí kho đi cùng suốt buổi kiểm tra (dẫn đường, lấy hàng, giải đáp), không được từ chối phối hợp'),
     ('仓储部主管（指派带路人）', 'Trưởng bộ phận kho (chỉ định người dẫn đường)'),
     ('检查过程中', 'Trong quá trình kiểm tra')),
    ('2',
     ('现场设备支持', 'Hỗ trợ thiết bị hiện trường'),
     ('配合提供扫码枪、叉车、登高设备；高位物料配合安全取用',
      'Phối hợp cung cấp máy quét, xe nâng, thiết bị leo cao; vật tư trên cao phối hợp lấy an toàn'),
     ('仓储部主管', 'Trưởng bộ phận kho'),
     ('检查过程中', 'Trong quá trình kiểm tra')),
    ('3',
     ('现场操作配合', 'Phối hợp thao tác hiện trường'),
     ('按检查要求开仓/开柜/移货；保管员在岗配合查阅系统、提供单据原件',
      'Mở kho/mở tủ/di chuyển hàng theo yêu cầu kiểm tra; thủ kho trực phối hợp tra cứu hệ thống, cung cấp chứng từ gốc'),
     ('当班保管员', 'Thủ kho trực ca'),
     ('检查过程中', 'Trong quá trình kiểm tra')),
    ('4',
     ('差异确认签字', 'Xác nhận chênh lệch và ký tên'),
     ('盘点结束当场与检查组长确认差异并双方签字；拒签记入报告「配合异常」',
      'Sau kiểm kê, xác nhận chênh lệch trực tiếp với tổ trưởng kiểm tra và hai bên ký tên; từ chối ký sẽ ghi vào báo cáo "phối hợp bất thường"'),
     ('仓库负责人 / 当班保管员', 'Người phụ trách kho / thủ kho trực ca'),
     ('盘点结束当场', 'Ngay khi kết thúc kiểm kê')),
    ('5',
     ('整改闭环', 'Khép kín khắc phục'),
     ('收到《整改事项清单》后48小时内反馈整改措施、责任人、承诺完成日；超期未完成将升级处理',
      'Trong vòng 48 giờ sau khi nhận 《Danh sách vấn đề cần khắc phục》 phản hồi biện pháp khắc phục, người chịu trách nhiệm, ngày cam kết hoàn thành; quá hạn chưa hoàn thành sẽ nâng cấp xử lý'),
     ('仓库负责人', 'Người phụ trách kho'),
     ('收到清单后48小时内', 'Trong 48 giờ sau khi nhận danh sách')),
]

table = doc.add_table(rows=1 + len(rows_data), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

widths = [Cm(1.2), Cm(3.2), Cm(6.8), Cm(3.0), Cm(2.8)]

# 表头
for ci, (zh, vi) in enumerate(headers):
    cell = table.rows[0].cells[ci]
    cell.width = widths[ci]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bilingual(p, zh, vi, size=10.5, bold=True)
    # 表头底纹
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'D9E2F3')
    tcPr.append(shd)

# 内容
for ri, row_data in enumerate(rows_data, start=1):
    for ci, val in enumerate(row_data):
        cell = table.rows[ri].cells[ci]
        cell.width = widths[ci]
        p = cell.paragraphs[0]
        if ci == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_bilingual(p, val, val, size=10.5)
        else:
            zh, vi = val
            add_bilingual(p, zh, vi, size=10.5)

# 固定布局
tbl = table._tbl
tblPr = tbl.tblPr
layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
tblPr.append(layout)

doc.save(OUT)
print('✅ 已生成:', OUT)
